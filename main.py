from fastapi import FastAPI, File, UploadFile, HTTPException, Form
from sqlalchemy import text
from database import engine
import pandas as pd  # Importa Pandas
from fastapi.middleware.cors import CORSMiddleware
from modelos import GetCliente, ResponseModel, SaveCliente, Vehiculo, GetOrden, GetVehiculo, SaveOrden, DatosLogin, Token, OrdenCompleta, Roles, Estatus, SaveUsuario, saveVehiculo, ImageData, Empleado,OrdenService,Checklist,CheckListHistorico,Flotillas
from fastapi.responses import JSONResponse
import json
from typing import List
from negocios import Negocios
from datetime import timedelta
from utils import utilsclass
import os
import pdfkit
import base64
from PIL import Image
from io import BytesIO
from fastapi.responses import StreamingResponse
from typing import Dict
from docx import Document
from docx.shared import Pt, Inches
from pydantic import BaseModel
from typing import Dict, List
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
ACCESS_TOKEN_EXPIRE_MINUTES = 480
import logging
from PIL import Image
import io
from sqlalchemy.sql import text
# Configuración del logger
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI()

origins = ["*"]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
options = {
    'page-size': 'A4',
    'margin-top': '1cm',
    'margin-right': '1cm',
    'margin-bottom': '1cm',
    'margin-left': '1cm',
}

class DocumentRequestV2(BaseModel):
    id_checklist: int  # Nuevo parámetro para identificar el checklist
    placeholders: Dict[str, str]
    logo_base64: str
    logo_derecho_base64: str

# Función para comprimir imágenes
def compress_image(image_path: str, quality: int = 85) -> bytes:
    img = Image.open(image_path)
    img = img.convert("RGB")  # Convertir a RGB si la imagen está en otro formato
    buffer = BytesIO()
    img.save(buffer, format="JPEG", quality=quality)  # Comprimir la imagen
    return buffer.getvalue()

# Función para validar el tamaño de una imagen Base64
def validate_image_size(base64_image: str, max_size_mb: int = 5) -> bool:
    # Recortar prefijo si está presente
    if base64_image.startswith("data:image/"):
        base64_image = base64_image.split(",")[1]
    
    # Calcular tamaño del contenido
    image_size = len(base64_image) * 3 / 4
    image_size_mb = image_size / (1024 * 1024)
    if image_size_mb > max_size_mb:
        raise ValueError(f"El tamaño de la imagen excede el límite de {max_size_mb} MB.")
    return True

# Función para obtener imágenes desde la base de datos para un checklist específico
def get_service_one(id_checklist: int, num_placeholders: int) -> List[str]:
    query = text("EXEC [dbo].[sp_get_all_checklist_Evidencias] @IdCheckList = :id_checklist")
    
    try:
        with engine.connect() as connection:
            result = connection.execute(query, {"id_checklist": id_checklist})
            columns = result.keys()
            rows = result.fetchall()
            roles_df = pd.DataFrame(rows, columns=columns)
    except Exception as e:
        logging.error(f"Error ejecutando el procedimiento almacenado: {e}")
        raise HTTPException(status_code=500, detail=f"Error ejecutando el procedimiento almacenado: {e}")
    
    if roles_df.empty:
        raise HTTPException(status_code=404, detail="No se encontraron datos para el checklist proporcionado.")
    
    # Filtrar columnas de imágenes
    image_columns = [col for col in roles_df.columns if isinstance(col, str) and '_foto' in col]
    
    if not image_columns:
        raise ValueError("El procedimiento almacenado no retornó columnas relacionadas con imágenes.")
    
    # Extraer imágenes de las columnas
    image_list = []
    for col in image_columns:
        image_list.extend(roles_df[col].dropna().tolist())
    
    # Filtrar imágenes válidas
    image_list = [img for img in image_list if img.strip() != '']
    
    # Asegurar que el número de imágenes no exceda los placeholders
    if len(image_list) > num_placeholders:
        image_list = image_list[:num_placeholders]
    
    image_list_base64 = []
    for img in image_list:
        try:
            validate_image_size(img)  # Validar tamaño de la imagen
            image_list_base64.append(img)  # La imagen ya está en formato Base64
        except Exception as e:
            logging.error(f"Error al procesar la imagen: {e}")
            raise ValueError(f"No se pudo procesar la imagen: {e}")
    
    return image_list_base64


# Función para generar el documento Word con imágenes y placeholders
def generate_word_documentv2(placeholders: Dict[str, str], images_base64: List[str], logo_base64: str, logo_derecho_base64: str) -> BytesIO:
    doc = Document()

    # Crear el encabezado
    section = doc.sections[-1]
    header = section.header
    header_table = header.add_table(rows=1, cols=3, width=Inches(6.5))
    header_table.autofit = True

    # Insertar logo izquierdo
    if logo_base64:
        image_data = base64.b64decode(logo_base64)
        image_stream = BytesIO(image_data)
        left_cell = header_table.cell(0, 0)
        paragraph = left_cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(image_stream, width=Inches(1.5))
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Insertar título centrado
    title_cell = header_table.cell(0, 1)
    title_paragraph = title_cell.paragraphs[0]
    set_header_format(title_paragraph, "FORMATO DE EVIDENCIAS FOTOGRÁFICAS")

    # Insertar logo derecho
    if logo_derecho_base64:
        image_data_right = base64.b64decode(logo_derecho_base64)
        image_stream_right = BytesIO(image_data_right)
        right_cell = header_table.cell(0, 2)
        paragraph_right = right_cell.paragraphs[0]
        run_right = paragraph_right.add_run()
        run_right.add_picture(image_stream_right, width=Inches(1.5))
        paragraph_right.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Espacio después del encabezado
    doc.add_paragraph()

    # Crear tabla de datos del vehículo
    table_data = doc.add_table(rows=3, cols=4)
    table_data.style = "Table Grid"

    keys = list(placeholders.keys())
    values = list(placeholders.values())

    # Ajustar keys_and_values en caso de que el número de placeholders sea menor que 6
    keys_and_values = [
        (keys[i], values[i]) if i < len(keys) else ("", "") for i in range(6)
    ]

    for i in range(3):  # Llenar filas y columnas
        for j in range(4):
            index = i * 4 + j
            if index < len(keys_and_values):
                key, value = keys_and_values[index]
                cell = table_data.cell(i, j)
                cell.text = f"{key.upper()}: {value}"
                cell.paragraphs[0].runs[0].font.size = Pt(10)

    # Espacio entre tabla y fotos
    doc.add_paragraph()

    # Crear tabla para imágenes (2 columnas)
    num_images = len(images_base64)
    rows = (num_images + 1) // 2 if num_images > 1 else 1  # Asegurar que haya al menos una fila

    table_images = doc.add_table(rows=rows, cols=2)
    table_images.style = "Table Grid"

    # Verificar que haya imágenes base64 antes de intentar insertarlas
    if not images_base64:
        raise ValueError("No se encontraron imágenes válidas para agregar.")

    for idx, image_base64 in enumerate(images_base64):
        image_data = base64.b64decode(image_base64)
        image_stream = BytesIO(image_data)

        row = idx // 2
        col = idx % 2
        cell = table_images.cell(row, col)

        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(image_stream, width=Inches(2.5), height=Inches(2.5))

    # Guardar documento en BytesIO
    word_stream = BytesIO()
    doc.save(word_stream)
    word_stream.seek(0)

    return word_stream

@app.post("/generate_and_downloadservice/")
async def generate_and_downloadv2(request: DocumentRequestV2):
    try:
        # Validar que el id_checklist sea un entero
        if not isinstance(request.id_checklist, int):
            raise HTTPException(status_code=400, detail="El id_checklist debe ser un entero.")
        
        # Calcular el número de placeholders desde la solicitud
        num_placeholders = len(request.placeholders)
        
        # Obtener las imágenes en base64 para el checklist proporcionado
        images_base64 = get_service_one(request.id_checklist, num_placeholders)
        logging.info(f"Imágenes obtenidas para id_checklist {request.id_checklist}: {images_base64}")
        
        if not images_base64:
            raise HTTPException(status_code=404, detail="No se encontraron imágenes para el checklist proporcionado.")
        
        # Generar el documento Word
        word_stream = generate_word_documentv2(
            request.placeholders,
            images_base64,
            request.logo_base64,
            request.logo_derecho_base64
        )

        # Devolver el documento como respuesta
        return StreamingResponse(
            word_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=EvidenciaFotografica.docx"}
        )
    
    except HTTPException as e:
        raise e
    except Exception as e:
        logging.error(f"Error inesperado: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error generando el documento: {str(e)}")
    
class DocumentRequest(BaseModel):
    placeholders: Dict[str, str]
    images_base64: List[str]
    logo_base64: str
    logo_derecho_base64: str  # Nuevo parámetro para el logo del lado derecho

def set_header_format(paragraph, text):
    """Establecer formato para el título del encabezado."""
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def generate_word_document(placeholders: Dict[str, str], images_base64: List[str], logo_base64: str, logo_derecho_base64: str) -> BytesIO:
    doc = Document()

    # Crear el encabezado
    section = doc.sections[-1]
    header = section.header
    header_table = header.add_table(rows=1, cols=3, width=Inches(6.5))
    header_table.autofit = True

    # Insertar logo izquierdo
    if logo_base64:
        image_data = base64.b64decode(logo_base64)
        image_stream = BytesIO(image_data)
        left_cell = header_table.cell(0, 0)
        paragraph = left_cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(image_stream, width=Inches(1.5))
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Insertar título centrado
    title_cell = header_table.cell(0, 1)
    title_paragraph = title_cell.paragraphs[0]
    set_header_format(title_paragraph, "FORMATO DE EVIDENCIAS FOTOGRÁFICAS")

    # Insertar logo derecho
    if logo_derecho_base64:
        image_data_right = base64.b64decode(logo_derecho_base64)
        image_stream_right = BytesIO(image_data_right)
        right_cell = header_table.cell(0, 2)
        paragraph_right = right_cell.paragraphs[0]
        run_right = paragraph_right.add_run()
        run_right.add_picture(image_stream_right, width=Inches(1.5))
        paragraph_right.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Espacio después del encabezado
    doc.add_paragraph()

    # Crear tabla de datos del vehículo
    table_data = doc.add_table(rows=0, cols=2)
    table_data.style = "Table Grid"

    for key, value in placeholders.items():
        row = table_data.add_row()
        row.cells[0].text = key.upper() + ":"
        row.cells[1].text = value
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)

    # Espacio entre tabla y fotos
    doc.add_paragraph()

    # Crear tabla para imágenes (2 columnas)
    num_images = len(images_base64)
    rows = (num_images + 1) // 2  # Calcular filas necesarias
    table_images = doc.add_table(rows=rows, cols=2)
    table_images.style = "Table Grid"

    for idx, image_base64 in enumerate(images_base64):
        image_data = base64.b64decode(image_base64)
        image_stream = BytesIO(image_data)

        row = idx // 2
        col = idx % 2
        cell = table_images.cell(row, col)

        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(image_stream, width=Inches(2.5), height=Inches(2.5))

    # Guardar documento en BytesIO
    word_stream = BytesIO()
    doc.save(word_stream)
    word_stream.seek(0)

    return word_stream

@app.post("/generate_and_download/")
async def generate_and_download(request: DocumentRequest):
    try:
        word_stream = generate_word_document(
            request.placeholders, request.images_base64, request.logo_base64, request.logo_derecho_base64
        )
        return StreamingResponse(word_stream,
                                 media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                 headers={"Content-Disposition": "attachment; filename=EvidenciaFotografica.docx"})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generando el documento: {str(e)}")


@app.post(
    path="/api/seguridad/iniciarsesion",
    name='Inicio de sesion',
    tags=['Seguridad'],
    description='Método para iniciar sesión',
    response_model=Token
)
async def login(payload: DatosLogin):
    _negocios = Negocios()
    user = await _negocios.getusuario(payload)
    if not user:
        return JSONResponse(
            status_code=401,
            content={"Id_Resultado": 0, "Respuesta": "Datos de acceso incorrectos"}
        )
    access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    access_token = await utilsclass.create_access_token(
        data={"sub": payload.usuario, "idUsuario": user["IdUsuario"], "idRol": user["Rol"]}, 
        expires_delta=access_token_expires
    )
    return {"access_token": access_token, "token_type": "bearer"}



# app.get(
#         path="/api/cliente",
#         name='Obtener clien',
#         tags=['Cliente'],
#         description='Método para obtener la informacion de un cliente',
#         response_model=GetCliente
# )
# def getcliente(IdCliente = int):
#     query = f"exec Clientes.clienteinsupdel @Accion = 4,@IdCliente ={IdCliente}"
#     roles_df = pd.read_sql(query, engine)
#     resultado = roles_df.to_dict(orient="records")
#     return JSONResponse(status_code=200,content=resultado)
# @app.get(
#         path="/api/clientes",
#         name='Obtener clientes',
#         tags=['Cliente'],
#         description='Método para obtener la informacion del cliente',
#         response_model=List[GetCliente]
# )
# def getclientes(busqueda = ""):
#     query = f"exec Clientes.clienteinsupdel @Accion = 2,@ParametroBusqueda = '{busqueda}' "
#     roles_df = pd.read_sql(query, engine)
#     resultado = roles_df.to_dict(orient="records")
#     return JSONResponse(status_code=200,content=resultado)

@app.get(
        path="/api/cliente",
        name='Obtener cliente',
        tags=['Cliente'],
        description='Método para obtener la informacion de un cliente',
        response_model=List[GetCliente]
)
def getcliente(idCliente = 0):
    query = f"exec Clientes.clienteinsupdel @Accion = 4,@IdCliente ={idCliente}"
    roles_df = pd.read_sql(query, engine)
    resultado = roles_df.to_dict(orient="records")
    resultado = resultado[0]
    return JSONResponse(status_code=200,content=resultado)

@app.get(
        path="/api/clientes",
        name='Obtener clientes',
        tags=['Cliente'],
        description='Método para obtener la informacion de todos los clientes',
        response_model=List[GetCliente]
)
def getclientes(busqueda = ""):
    query = f"exec Clientes.clienteinsupdel @Accion = 2,@ParametroBusqueda = '{busqueda}' "
    roles_df = pd.read_sql(query, engine)
    resultado = roles_df.to_dict(orient="records")
    return JSONResponse(status_code=200,content=resultado)


@app.post(
        path="/api/cliente",
        name='Guardar cliente',
        tags=['Cliente'],
        description='Método para guardar la informacion del cliente}',
        response_model=ResponseModel
)
def saveCliente(payload: SaveCliente ):
    query = f"EXEC clientes.clienteinsupdel \
    @Nombre='{payload.Nombre}', \
    @Calle='{payload.Calle}',  \
    @Colonia='{payload.Colonia}', \
    @Ciudad='{payload.Ciudad}', \
    @Estado='{payload.Estado}', \
    @Tel='{payload.Tel}', \
    @Cel='{payload.Cel}', \
    @Email='{payload.Email}', \
    @RFC='{payload.RFC}', \
    @Autorizacion_ext='{payload.Autorizacion_ext}', \
    @No_int='{payload.No_int}', \
    @Facturar_a='{payload.Facturar_a}', \
    @IdUsuarioEmpleado='{payload.Id_empleado}', @Accion = 1" 
    print(query)

    with engine.begin() as conn:
        conn.execution_options(autocommit = True)
        roles_df = pd.read_sql(query, conn)
    dumpp = ResponseModel(id_resultado=1,respuesta="El cliente se guardo de manera correcta")
    dict = dumpp.model_dump()
    return JSONResponse(status_code=200, content=dict)
#################################################################################################################
@app.get(
    path="/api/vehiculo",
    name='Obtener vehiculo',
    tags=['Vehiculo'],
    description='Método para obtener la información de un vehículo',
    response_model=GetVehiculo
)
def getvehiculo(idVehiculo=0):
    query = f"exec [dbo].[ObtenerVehiculo] @IdVehiculo = {idVehiculo}"
    vehiculo_df = pd.read_sql(query, engine)
    resultado = vehiculo_df.to_dict(orient="records")
    
    # Verifica si el resultado no está vacío antes de devolver
    if resultado:
        return JSONResponse(status_code=200, content=resultado[0])
    else:
        return JSONResponse(status_code=404, content={"message": "Vehículo no encontrado."})

@app.get(
    path="/api/vehiculos",
    name='Obtener vehiculos',
    tags=['Vehiculo'],
    description='Método para obtener la información de todos los vehículos',
    response_model=List[GetVehiculo]
)
def getvehiculos(parametro=""):
    query = f"exec [dbo].[ObtenerVehiculo] @ParametroBusqueda = '{parametro}'"
    vehiculos_df = pd.read_sql(query, engine)
    resultado = vehiculos_df.to_dict(orient="records")
    
    # Devuelve la lista de vehículos
    return JSONResponse(status_code=200, content=resultado)
#################################################################################################################
# @app.get(
#         path="/api/vehiculo",
#         name='Obtener vehiculo',
#         tags=['Vehiculo'],
#         description='Método para obtener la informacion de todos los vehiculos',
#         response_model=List[GetVehiculo]
# )
# def getvehiculos(idVehiculo = 0):
#     query = f"exec [dbo].[ObtenerVehiculo] @IdVehiculo = {idVehiculo}"
#     roles_df = pd.read_sql(query, engine)
#     resultado = roles_df.to_dict(orient="records")
#     return JSONResponse(status_code=200,content=resultado[0])


# @app.get(
#     path="/api/vehiculos",
#         name='Obtener vehiculos',
#         tags=['Vehiculo'],
#         description='Método para obtener la informacion de un vehiculo',
#         response_model=GetVehiculo
# )
# def getvehiculos(parametro = ""):
#     query = f"exec [dbo].[ObtenerVehiculo] @ParametroBusqueda = '{parametro}'"
#     roles_df = pd.read_sql(query, engine)
#     resultado = roles_df.to_dict(orient="records")
#     return JSONResponse(status_code=200,content=resultado)

@app.post(
        path="/api/vehiculo",
        name='Guarda vehiculo',
        tags=['Vehiculo'],
        description='Método para guardar la informacion de los vehiculos del cliente}',
    response_model=ResponseModel
)
def guardarVehiculo(payload: saveVehiculo):
    try:
        # Convertir listas de fotos a cadenas de Base64 separadas por comas
        fotos = {
            "MotorVehiculo_foto": ",".join(payload.MotorVehiculo_foto),
            "Acumulador_foto": ",".join(payload.Acumulador_foto),
            "Espejo_retrovisor_foto": ",".join(payload.Espejo_retrovisor_foto),
            "Espejo_izquierdo_foto": ",".join(payload.Espejo_izquierdo_foto),
            "Espejo_derecho_foto": ",".join(payload.Espejo_derecho_foto),
            "Antena_foto": ",".join(payload.Antena_foto),
            "Tapones_ruedas_foto": ",".join(payload.Tapones_ruedas_foto),
            "Radio_foto": ",".join(payload.Radio_foto),
            "Encendedor_foto": ",".join(payload.Encendedor_foto),
            "Gato_foto": ",".join(payload.Gato_foto),
            "Herramienta_foto": ",".join(payload.Herramienta_foto),
            "Llanta_refaccion_foto": ",".join(payload.Llanta_refaccion_foto),
            "Limpiadores_foto": ",".join(payload.Limpiadores_foto),
            "Pintura_rayada_foto": ",".join(payload.Pintura_rayada_foto),
            "Cristales_rotos_foto": ",".join(payload.Cristales_rotos_foto),
            "Golpes_foto": ",".join(payload.Golpes_foto),
            "Tapetes_foto": ",".join(payload.Tapetes_foto),
            "Extintor_foto": ",".join(payload.Extintor_foto),
            "Tapones_gasolina_foto": ",".join(payload.Tapones_gasolina_foto),
            "Calaveras_rotas_foto": ",".join(payload.Calaveras_rotas_foto),
            "Molduras_completas_foto": ",".join(payload.Molduras_completas_foto),
            "MotorVehiculo_video": ",".join(payload.MotorVehiculo_video),
            "Acumulador_video": ",".join(payload.Acumulador_video),
            "Espejo_retrovisor_video": ",".join(payload.Espejo_retrovisor_video),
            "Espejo_izquierdo_video": ",".join(payload.Espejo_izquierdo_video),
            "Espejo_derecho_video": ",".join(payload.Espejo_derecho_video),
            "Antena_video": ",".join(payload.Antena_video),
            "Tapones_ruedas_video": ",".join(payload.Tapones_ruedas_video),
            "Radio_video": ",".join(payload.Radio_video),
            "Encendedor_video": ",".join(payload.Encendedor_video),
            "Gato_video": ",".join(payload.Gato_video),
            "Herramienta_video": ",".join(payload.Herramienta_video),
            "Llanta_refaccion_video": ",".join(payload.Llanta_refaccion_video),
            "Limpiadores_video": ",".join(payload.Limpiadores_video),
            "Pintura_rayada_video": ",".join(payload.Pintura_rayada_video),
            "Cristales_rotos_video": ",".join(payload.Cristales_rotos_video),
            "Golpes_video": ",".join(payload.Golpes_video),
            "Tapetes_video": ",".join(payload.Tapetes_video),
            "Extintor_video": ",".join(payload.Extintor_video),
            "Tapones_gasolina_video": ",".join(payload.Tapones_gasolina_video),
            "Calaveras_rotas_video": ",".join(payload.Calaveras_rotas_video),
            "Molduras_completas_video": ",".join(payload.Molduras_completas_video)
           
        }
        
        # Crear el diccionario de parámetros sin conflicto
        parametros = payload.dict(exclude=fotos.keys())
        parametros.update(fotos)

        query = text("""
            exec dbo.InsertarVehiculo 
                @Id_Cliente = :IdCliente,
                @Id_Empleado = :Id_empleado,
                @Marca = :Marca,
                @Modelo = :Modelo,
                @Color = :Color,
                @No_serie = :No_serie,
                @Placa = :Placa,
                @Tipo = :Tipo,
                @Motor = :Motor,
                @Kms = :Kms,
                @MotorVehiculo = :MotorVehiculo,
                @Acumulador = :Acumulador,
                @Espejo_retrovisor = :Espejo_retrovisor,
                @Espejo_izquierdo = :Espejo_izquierdo,
                @Espejo_derecho = :Espejo_derecho,
                @Antena = :Antena,
                @Tapones_ruedas = :Tapones_ruedas,
                @Radio = :Radio,
                @Encendedor = :Encendedor,
                @Gato = :Gato,
                @Herramienta = :Herramienta,
                @Llanta_refaccion = :Llanta_refaccion,
                @Limpiadores = :Limpiadores,
                @Pintura_rayada = :Pintura_rayada,
                @Cristales_rotos = :Cristales_rotos,
                @Golpes = :Golpes,
                @Tapetes = :Tapetes,
                @Extintor = :Extintor,
                @Tapones_gasolina = :Tapones_gasolina,
                @Calaveras_rotas = :Calaveras_rotas,
                @Molduras_completas = :Molduras_completas,
                @MotorVehiculo_foto = :MotorVehiculo_foto,
                @Acumulador_foto = :Acumulador_foto,
                @Espejo_retrovisor_foto = :Espejo_retrovisor_foto,
                @Espejo_izquierdo_foto = :Espejo_izquierdo_foto,
                @Espejo_derecho_foto = :Espejo_derecho_foto,
                @Antena_foto = :Antena_foto,
                @Tapones_ruedas_foto = :Tapones_ruedas_foto,
                @Radio_foto = :Radio_foto,
                @Encendedor_foto = :Encendedor_foto,
                @Gato_foto = :Gato_foto,
                @Herramienta_foto = :Herramienta_foto,
                @Llanta_refaccion_foto = :Llanta_refaccion_foto,
                @Limpiadores_foto = :Limpiadores_foto,
                @Pintura_rayada_foto = :Pintura_rayada_foto,
                @Cristales_rotos_foto = :Cristales_rotos_foto,
                @Golpes_foto = :Golpes_foto,
                @Tapetes_foto = :Tapetes_foto,
                @Extintor_foto = :Extintor_foto,
                @Tapones_gasolina_foto = :Tapones_gasolina_foto,
                @Calaveras_rotas_foto = :Calaveras_rotas_foto,
                @Molduras_completas_foto = :Molduras_completas_foto,
                @MotorVehiculo_video = :MotorVehiculo_video,
                @Acumulador_video = :Acumulador_video,
                @Espejo_retrovisor_video = :Espejo_retrovisor_video,
                @Espejo_izquierdo_video = :Espejo_izquierdo_video,
                @Espejo_derecho_video = :Espejo_derecho_video,
                @Antena_video = :Antena_video,
                @Tapones_ruedas_video = :Tapones_ruedas_video,
                @Radio_video = :Radio_video,
                @Encendedor_video = :Encendedor_video,
                @Gato_video = :Gato_video,
                @Herramienta_video = :Herramienta_video,
                @Llanta_refaccion_video = :Llanta_refaccion_video,
                @Limpiadores_video = :Limpiadores_video,
                @Pintura_rayada_video = :Pintura_rayada_video,
                @Cristales_rotos_video = :Cristales_rotos_video,
                @Golpes_video = :Golpes_video,
                @Tapetes_video = :Tapetes_video,
                @Extintor_video = :Extintor_video,
                @Tapones_gasolina_video = :Tapones_gasolina_video,
                @Calaveras_rotas_video = :Calaveras_rotas_video,
                @Molduras_completas_video = :Molduras_completas_video,
                @IdFlotilla = :IdFlotilla
                
        """)

        # Ejecutar la consulta pasando `parametros` como un solo diccionario
        with engine.begin() as conn:
            conn.execute(query, parametros)

        # Respuesta de éxito
        return JSONResponse(status_code=200, content={
            "id_resultado": 1,
            "respuesta": "Se guardó la información del vehículo de manera correcta",
            "detalles": parametros
        })

    except Exception as e:
        # Respuesta de error
        raise HTTPException(status_code=500, detail=f"Error al guardar el vehículo: {str(e)}")
@app.put(
        path="/api/vehiculo",
        name='Actualizar vehiculo',
        tags=['Vehiculo'],
        description='Método para actualizar la informacion de los vehiculos del cliente}',
    response_model=ResponseModel
)
def updateVehiculo(payload: GetVehiculo):
    dumpp = ResponseModel(id_resultado=1,respuesta="Se actualizó la información del vehiculo de manera correcta")
    dict = dumpp.model_dump()
    return JSONResponse(status_code=200, content=dict)

@app.put(
        path="/api/cliente",
        name='Actualizar cliente',
        tags=['Cliente'],
        description='Método para actualizar la informacion del cliente}',
        response_model=ResponseModel
)
def putcliente(payload: GetCliente ):
    query = f"exec [Clientes].[clienteinsupdel] @Accion = 3, @idCliente = {payload.ID}, @Nombre = '{payload.Nombre}', @Calle = '{payload.Calle}' \
        ,@Colonia = '{payload.Colonia}', @Ciudad = '{payload.Ciudad}',  @Estado = '{payload.Estado}', @Tel = '{payload.Tel}', @Cel = '{payload.Cel}' \
        ,@Email = '{payload.Email}', @RFC = '{payload.RFC}', @Autorizacion_ext = '{payload.Autorizacion_ext}', @No_int = '{payload.No_int}',@Facturar_a = '{payload.Facturar_a}' \
        ,@IdUsuarioEmpleado = '{payload.Id_empleado}'"
    with engine.begin() as conn:
        conn.execution_options(autocommit = True)
        roles_df = pd.read_sql(query, conn)
    dumpp = ResponseModel(id_resultado=1,respuesta="El cliente se actualizo de manera correcta")
    dict = dumpp.model_dump()
    return JSONResponse(status_code=200, content=dict)

@app.post(
    path="/api/ordenserviciofull",
        name='Guardar orden}',
        tags=['Orden'],
        description='Método para guardan la orden}',
        response_model=ResponseModel
)
def saveorden(payload: OrdenCompleta):
    
    dumpp = ResponseModel(id_resultado=1,respuesta="Se guardo la orden")
    dict = dumpp.model_dump()
    return JSONResponse(status_code=200, content=dict)

@app.get(
        path="/api/roles",
        name='Obtener roles',
        tags=['Catalogos'],
        description='Método para obtener los roles}',
        response_model=Roles
)
def obtener_roles():
    query = "SELECT * FROM Catalogos.Roles"
    roles_df = pd.read_sql(query, engine)
    resultado = roles_df.to_dict(orient="records")
    return JSONResponse(status_code=200,content=resultado)

@app.get(
        path="/api/estatus",
        name='Obtener estatus',
        tags=['Catalogos'],
        description='Método para obtener los estatus}',
    response_model=Estatus
)
def obtener_roles():
    query = "SELECT * FROM Catalogos.Estatus"
    roles_df = pd.read_sql(query, engine)
    resultado = roles_df.to_dict(orient="records")
    return JSONResponse(status_code=200,content=resultado)

@app.post(
    path="/api/usuarios",
    name='crear usuario',
    tags=['Seguridad'],
    description='Método para crear usuarios',
    response_model=ResponseModel
)
async def crearusuario(payload: SaveUsuario):
    passhash = await utilsclass.get_password_hash(payload.Password)
    query = f"exec Seguridad.usuariosinsupdel @Nombre = '{payload.Nombre}', @Password = '{passhash}', @Rol = {payload.Rol}, @Estatus = {payload.Estatus}, @Accion = 1"
    with engine.begin() as conn:
        conn.execution_options(autocommit = True)
        roles_df = pd.read_sql(query, conn)
    dumpp = ResponseModel(id_resultado=1,respuesta="Se agregó el usuario de manera correcta")
    dict = dumpp.model_dump()
    return JSONResponse(status_code=200, content=dict)

@app.post(
    path="/api/ordenserviciopdf",
    name='obtener pdf de la orden de servicio',
    tags=['Orden'],
    description='Método para obtener el pdf de la orden de servicio',
    response_model=ResponseModel
)
def convert_html_to_pdf(clienteId: int):
    try:
        query = f"exec [Clientes].[ordendeservicio]  @idCliente = {clienteId}"
        with engine.begin() as conn:
            conn.execution_options(autocommit=True)
            roles_df = pd.read_sql(query, conn)

        # Asegúrate de que hay al menos dos filas en el DataFrame
        
            # Accede al valor en la segunda fila de la columna 'Nombre'
            

            # Construye el HTML con el valor obtenido
        orden = roles_df['idOrden'].iloc[0]
        nombre = roles_df['Nombre'].iloc[0]
        calle =roles_df['Calle'].iloc[0]
        factura =roles_df['Facturar_a'].iloc[0]
        noint = roles_df['No_int'].iloc[0]
        colonia = roles_df['Colonia'].iloc[0]
        ciuadad = roles_df['Ciudad'].iloc[0]
        estado = roles_df['Estado'].iloc[0]
        tel = roles_df['Tel'].iloc[0]
        cel = roles_df['Cel'].iloc[0]
        email = roles_df['Email'].iloc[0]
        rfc = roles_df['RFC'].iloc[0]
        autorizacion = roles_df['Autorizacion_ext'].iloc[0]
        marca = roles_df['Marca'].iloc[0]
        tipo = roles_df['Tipo'].iloc[0] 
        modelo = roles_df['Modelo'].iloc[0]
        motor = roles_df['Motor'].iloc[0]
        color = roles_df['Color'].iloc[0]
        kms = roles_df['kms'].iloc[0]
        noserie = roles_df['No_Serie'].iloc[0]
        placa = roles_df['Placa'].iloc[0]
        espejoretrovisor = roles_df['Espejo_retrovisor'].iloc[0]
        espejoizq = roles_df['Espejo_izquierdo'].iloc[0]
        antena = roles_df['Antena'].iloc[0]
        taponesruedas = roles_df['Tapones_ruedas'].iloc[0]
        radio = roles_df['Radio'].iloc[0]
        encendedor = roles_df['Encendedor'].iloc[0]
        gato = roles_df['Gato'].iloc[0]
        herramienta = roles_df['Herramienta'].iloc[0]
        llanarefaccion = roles_df['Llanta_refaccion'].iloc[0]
        limpiadores = roles_df['Limpiadores'].iloc[0]
        pintura = roles_df['Pintura_rayada'].iloc[0]
        cristales = roles_df['Cristales_rotos'].iloc[0]
        golpes = roles_df['Golpes'].iloc[0]
        tapetes = roles_df['Tapetes'].iloc[0]
        extintor = roles_df['Extintor'].iloc[0]
        tapones_gasolina = roles_df['Tapones_gasolina'].iloc[0]
        calaveras = roles_df['Calaveras_rotas'].iloc[0]
        molduras = roles_df['Molduras_completas'].iloc[0] 


 
        htmlstring = r"""<!DOCTYPE html>
        <html>
        <head>
        <meta charset="UTF-8">
        <div class="container">
        
        <div class="center-text">
        <strong> <p>ORDEN DE SERVICIO<br/>
        Refaccionaría y Taller Aautomotriz <br/>
        Trotamundos <br/>
        Horario <br/>
        Lunes a Viernes <br/>
        8 am - 6pm <br/>
        Sábados<br/>
        8 am - 3 pm<br/>        
        29 Guerrero y Bravo #422.
        Col. Héroes de Nacozari. C.P. 87030 <br/>
        Tel:(834) 2285 2869 y (834) 285 2872 R.F.C. CALN810623UF3</p>
        </strong>
        </div>
        
        </div>
  
		<!-- <img style="width:30%;" align:left; src="C:\Users\arael\Downloads\trot.jpg"/>
		<img style="width:30%;" align:right; src="C:\Users\arael\Downloads\todo.jpg"/> -->
		<!-- <h2 style="text-align:center;font-size:25px;">ORDEN DE SERVICIO</h2> 
		<h3 style="text-align:center;">Refaccionaría y Taller Aautomotriz</h3>
		<div style="text-align: right;">Horario <br/></h3>
		Lunes a Viernes <br/>
		8 am - 6pm <br/>
		Sábados<br/>
		8 am - 3 pm<br/>
		<h3 style="text-align:center;">Trotamundos</h3>
		<h3 style="text-align:center;">29 Guerrero y Bravo #422.</h3>
		<h3 style="text-align:center;">Col. Héroes de Nacozari. C.P. 87030</h3>
		<h3 style="text-align:center;">Tel:(834) 2285 2869 y (834) 285 2872 R.F.C. CALN810623UF3</h3> -->
		<div style="position: relative;">
		<div style="width: 20%; position: relative;">"""
        htmlstring +=f"""
		<table> 
		<th>
		  ORDEN
		</th>
		<tr>
		<td>
        {roles_df['idOrden'].iloc[0]}
		</td>
		</tr>
		<tr>
		<td>
        Fecha de Recepción:
	    </td>
		</tr>
		<tr>
		<td>
        Fecha de Entrega:
		</td>
		</tr>
		</table>
		</div>
		<div style="width: 100%; position: absolute; left: 30%; top: 0; height: 100%;">
		<table> 
		<tr>
        <td>Hora de compromiso</td>
        <td>Motivo de Visita:  (Previsto)  (Correctivo)</td>
		</tr>
		<tr>
        <td>Hora de entrega</td>
        <td>Medio:  Periódico  Radio  TV  Volante   Recomendación  Otros</td>
		</tr>
		</table>
		</div>
		</div>
		</head>
		<meta charset="UTF-8">"""
        htmlstring +="""
		<style>
		.left-image, .right-image {
        max-width: 20%; 
        height: 35%;
        border-radius:50%;
        }
		.container {
        display: flex;
        justify-content: space-between;
        align-items: center;
        width: 100%;
        }
        .left-image, .right-image {
        max-width: 20%; /* Puedes ajustar el tamaño de las imágenes según tus necesidades */
        height: auto;
        }
        .center-text {
        flex: 1;
        text-align: center;
        font-size: 20px;
        padding: 0 20px; /* Espaciado alrededor del texto */
        font-weight: 50%;
        }
		table {
        font-family: arial, sans-serif;
        border-collapse: collapse;
        width: 100%;
        border: 1px solid #000;
		}
		td, th {
        border-collapse: collapse; border: 1px solid #dddddd;
        text-align: left;
        border: 1px solid #000;
        vertical-align: top;
        padding: 8px;
		}
		th {
        background-color: #dddddd;
		}
		.linea {
		border-top: 1px solid black;
		height: 2px;
		max-width: 200px;
		padding: 0;
		margin: 20px auto 0 auto;
		}
		td.empty-cell {
        min-width: 50px; /* Establece el tamaño mínimo para celdas vacías */
        min-height: 20px; /* Establece la altura mínima para celdas vacías */
		}
		</style>"""
        htmlstring += f"""
		<body>
		<div style="position: relative;">
        <div style="width: 40%; position: relative; height: 100%;">
        <table>
        <th>Cliente</th>  
        <tr>
        <td>Facturar a:{factura}</td>            
        </tr>
        <tr>
        <td>Nombre:{nombre}</td> 
        </tr>
        <tr>
        <td>Calle:{calle} </td>
        <td>No. int:{noint}</td>
        </tr>
        <tr>
        <td>Colonia:{colonia}</td>
        <td>Ciudad:{ciuadad}</td>
        </tr>
        <tr>
        <td>Cumpleaños:</td>
        <td>Estado:{estado}</td>  
        </tr>
        <tr>
        <td>Tel:{tel}</td>
        <td>Cel. /Nex:{cel} </td>
        </tr>
        <tr>
        <td>EMAIL:{email}</td>
        </tr>
        <tr>
        <td>RFC:{rfc}</td> 
        </tr>
        <tr>
        <td>Autorización Externa:{autorizacion}</td>
        </tr>
        <tr>
        <th>Servicio</th>
        </tr>

        <tr>
        
        <td>
        <h3 style="text-align:center;">FILTROS</h3>
        </td>
        </tr>
        
        <tr>
          <td>Filtro de Aire:</td>
         
          
        </tr>
        <tr>
          <td>Filtro de Aceite:</td>
         
        </tr>
        
        <tr>
          <td>Filtro de Cabina</td>
         
        </tr>
        <tr>
          <td>Filtro de Gasolina:</td>
         
          
        </tr>
        <tr>
        </tr>
         <tr>
        
        <td>
        <h3 style="text-align:center;">LUBRICACIÓN</h3>
        </td>
        </tr>
        <tr>
          <td>Tapón de Carter:</td>
        
        </tr>
        <tr>
          <td>Aceite de Motor:</td> 
          <td>W  /SAE</td>
        
        
        </tr>
        <tr>
          <td>Aceite de Transmisión Automática:</td>
    
        </tr>
          <tr>
          <td>Aceite de Transmisión Estándar:</td>
          <td>W  /SAE</td>
        </tr>
          <tr>
          <td>Aceite de Diferencial:</td>
          <td>W  /SAE</td>
        </tr>
          <tr>
          <td>Dirección Hidráulica:</td>
        
        </tr>
          <tr>
          <td>Líquido de Frenos:</td>
          <td>W  /SAE</td>
       
        </tr>
          <tr>
          <td>Líquido Limpiaparabrisas:</td>
          <td>W  /SAE</td>
        
        </tr>
          <tr>
          <td>Aditivo:</td>
        
        </tr>
           <tr>
        
        <td>
        <h3 style="text-align:center;">BANDAS</h3>
        </td>
        </tr>
          <tr>
          <td>Poly-V:</td>
        
        </tr>
          <tr>
          <td>Alternador:</td>
         
        </tr>
          <tr>
          <td>Aire Acondicionado:</td>
        
        </tr>
            <tr>
          <td>Dirección Hidráulica:</td>
         
        </tr>
            <tr>
          <td>Banda de Tiempo:</td>
        
        </tr>

        <tr>        
        <td>
        <h3 style="text-align:center;">ENFRIAMIENTO</h3>
        </td>
        </tr>
         <tr>
          <td>Manguera Superior Radiador:</td>
         
        </tr>
           <tr>
          <td>Refrigerante/ Anticongelante:</td>
        
        </tr>
        
           <tr>
          <td>Bomba de Agua:</td>
         
        </tr>                  
           <tr>
          <td>Manguera Inferior Radiador:</td>
          
        </tr>
        <tr>
          <td>Tapón de Radiador:</td>
          
        </tr>

        <tr>
          
          <td>
          <h3 style="text-align:center;">AMORTIGUADORES/BASES</h3>
          </td>
         
          </tr>
          
        <tr>
            <td>Delanteros:</td>
          
          </tr>
          <tr>
            <td>Traseros:</td>
            
          </tr>
        <tr>
        
        <td>
        <h3 style="text-align:center;">MOTOR</h3>
        </td>
        
        
        </tr>
        
        <tr>
          <td>Diagnóstico (Costo o revisión):</td>
        
        </tr>
        
        <tr>
          <td>Afinación:</td>
          <td>Menor 	Mayor 	Premium</td>
        </tr>
        
        
        <tr>
          <td>Tapa de Distribuidor:</td>
         
        </tr>
        
        
        <tr>
          <td>Rotor de Distribuidor:</td>
        </tr>              
        <tr>
          <td>Cables de Bujias:</td>
        
        </tr>
        
        <tr>
          <td>Bujias:</td>
          
        </tr>                
        <tr>
          <td>Sensor:</td>
         
        </tr>
             
        </table>
        </div>
        <div style="width: 50%; position: absolute; left: 50%; top: 0; height: 100%;">
        <table>
        <th>Vehículo</th>
        <tr>                
        <td>Marca:{marca}</td> 
        <td>Tipo:{tipo}</td>  
        </tr>
        <tr>
        <td>Modelo:{modelo}</td>
        <td>Motor:{motor}</td>
        </tr>
        <tr>          
        <td>Color:{color}</td>
        <td>Kms:{kms}</td>
        </tr>
        <tr>
        
        <td>N. Serie:{noserie}</td>
          
        </tr>
        <tr>
          <td>Placa:{placa}</td>
        </tr>
        <tr>
          <th>Inventario de Vehículo</th>
        </tr>


        <tr>
        
        
        <td>Espejo Retrovisor:</td>
        <td>{espejoretrovisor}</td>
        
        
        </tr>
        <tr>
        
        <td>Espejo Izquierdo:</td>
        <td>{espejoizq}</td>
        </tr>
          
        <tr>
            
        <td>Espejo Derecho:</td>
        <td></td>
        </tr>
        <tr>
           
        <td>Antena:</td>
        <td>{antena}</td>
                  
        </tr>
        <tr>
                  
        <td>Tapones de Ruedas:</td>
        <td>{taponesruedas}</td>
        </tr>
        <tr>
         
        
        <tr>
                  
        <td>Radio:</td>
        <td>{radio}</td>
        </tr>
        <tr>
                   
        <td>Encendedor:</td>
        <td>{encendedor}</td>
        </tr>
        <tr>
                   
        <td>Gato:</td>
        <td>{gato}</td>
        </tr>
        <tr>
                   
        <td>Herramienta:</td>
        <td>{herramienta}</td>
        </tr>
        <tr>
                    
        <td>Llanta de Refacción:</td>
        <td>{llanarefaccion}</td>
        </tr>
        <tr>
        
        <td>Limpiadores:</td>
        <td>{limpiadores}</td>
        </tr>
        <tr>
                   
        <td>Pintura Rayada:</td>
        <td>{pintura}</td>
        </tr>
        <tr>
                    
        <td>Cristales Rotos:</td>
        <td>{cristales}</td>
        </tr>
        <tr>
                    
        <td>Golpes:</td>
        <td>{golpes}</td>
        </tr>

        <tr>
                    
        <td>Tapetes:</td>
        <td>{tapetes}</td>
        </tr>

        <tr>
          
        <td>Extintor:</td>
        <td>{extintor}</td>
        </tr>

        <tr>
          
        <td>Tapón de Gasolina:</td>
        <td>{tapones_gasolina}</td>
        </tr>

        <tr>
          
        <td>Calaveras Rotas:</td>
        <td>{calaveras}</td>
        </tr>

        <tr>
          
        <td>Molduras Completas:</td>
        <td>{molduras}</td>
        </tr>

        <tr>
        <th>
        Servicio
        </th>
        </tr>
        <tr> 
        <td>
        <h3 style="text-align:center;">LLANTAS</h3>
        </td>
        </tr>

        <tr>
        <td>Medida:</td>
        </tr>
        <tr>
        <td>Alineación:</td>
        </tr>
        <tr>
        <td>Balanceo:</td>
        </tr>
        <tr>
        <td>Rotación:</td>
        </tr>
        <tr>
        <td>Montaje:</td>
        </tr>
        <tr>
        <td>Válvulas:</td>
        </tr>
        <tr> 
        <td>
        <h3 style="text-align:center;">FRENOS</h3>
        </td>
        </tr>
        <tr>
        <td>Generales:</td>
        </tr>
        <tr>
        <td>Delanteros:</td>
        </tr>
        <tr>
        <td>Traseros:</td>
        </tr>
        <tr>
        <td>Discos (2) (4):</td>
        </tr>
        <tr>
        <td>Tambores:</td>
        </tr>
        <tr>
        <td>Cilindros:</td>
        </tr>
        <tr>
        <td>Limpieza y Ajuste:</td>
        </tr>
        <tr>
        <td>Rectificado Disc/Tam:</td>
        </tr>
        <tr> 
        <td>
        <h3 style="text-align:center;">OTROS</h3>
        </td>
        </tr>
        <tr>
        <td>Limpiaparabrisas:</td>
        </tr>
        <tr>
        <td>Engrasado:</td>
        </tr>
        <tr>
        <td>Químicos:</td>
        </tr>



        </table>
        </div>
		</div>

		<table>
		<th>
		Observaciones
		</th>
		<tr>
        <td>

        </td>
	    </tr>
		</table>
      

		<div></div>
      
		<table> 
         
      
		<tr>
		<td>
		Firma del Proveedor
      
		</td>

		<td>
		Firma de Aceptación del Cliente
		
		</td>
		
		</tr>
		
		</table>
		
		<div></div>

		<table>
		<th>
		Servicio Solicitado
		</th>
		<th>
		Recibió
		</th>
		<th>
		Técnico
		</th>
		<th>
		Orden
		</th>


		<tr>
		  <td></td>
		</tr>

		</table>
		</body>
		</html>"""
         # Resto del código para convertir a PDF...
        img = "\\img1.jpg"
        pdf_path = "example.pdf"
          # Rutas y configuraciones para Linux
        path_wkhtmltopdf = '/usr/local/bin/wkhtmltopdf'
        config = pdfkit.configuration(wkhtmltopdf=path_wkhtmltopdf)
        pdfkit.from_string(htmlstring, 'reporte.pdf', configuration=config)
        return JSONResponse(content={"message": "PDF creado exitosamente"}, status_code=200)
    except Exception as e:
        return JSONResponse(content={"error": str(e)}, status_code=500)

@app.get(
        path="/api/empleados",
        name='Obtener empleados',
        tags=['Seguridad'],
        description='Método para obtener la informacion de todos los empleados',
        response_model=List[Empleado]
)
def getempleados():
    query = f"[Seguridad].[usuariosinsupdel] @Accion = 2"
    roles_df = pd.read_sql(query, engine)
    resultado = roles_df.to_dict(orient="records")
    return JSONResponse(status_code=200,content=resultado)

@app.get(
        path="/api/empleado",
        name='Obtener empleado',
        tags=['Seguridad'],
        description='Método para obtener la informacion de todos los empleados',
        response_model=Empleado
)
def getempleados(IdUsuario: str):
    query = f"[Seguridad].[usuariosinsupdel] @Accion = 3, @Idusuario = {IdUsuario}"
    roles_df = pd.read_sql(query, engine)
    resultado = roles_df.to_dict(orient="records")
    return JSONResponse(status_code=200,content=resultado)

@app.post(
        path="/api/checklisthistorico",
        name='Insertar checklist histórico',
        tags=['ChecklistHistorico'],
        description='Método para insertar el histórico del checklist',
        response_model=CheckListHistorico
)
def savechecklisthistorico(payload: CheckListHistorico):
    query = f"""
    EXEC InsertarHistoricoCheckList
        @IdChecklist = {payload.IdChecklist}, \
        @IdVehiculo = {payload.IdVehiculo}, \
        @IdEmpleado = {payload.IdEmpleado}, \
        @Fecha = '{payload.Fecha}', \
        @TiempoTranscurrido = {payload.TiempoTranscurrido}, \
        @Estado = '{payload.Estado}', \
       """
    print (query)
    with engine.begin() as conn:
          conn.execution_options(autocommit = True)
          roles_df = pd.read_sql(query, conn)
    dumpp = ResponseModel(id_resultado=1,respuesta="El histórico del checklist se guardó de manera correcta")
    dict = dumpp.model_dump()
    return JSONResponse(status_code=200, content=dict)

@app.post(
        path="/api/checklisthistoricoservicio",
        name='Insertar checklist histórico servicio',
        tags=['ChecklistHistorico'],
        description='Método para insertar el histórico del checklist servicio',
        response_model=CheckListHistorico
)
def savechecklisthistoricoservicio(payload: CheckListHistorico):
    query = f"""
    EXEC InsertarHistoricoCheckListServicio
        @IdChecklist = {payload.IdChecklist}, \
        @IdVehiculo = {payload.IdVehiculo}, \
        @IdEmpleado = {payload.IdEmpleado}, \
        @Fecha = '{payload.Fecha}', \
        @TiempoTranscurrido = {payload.TiempoTranscurrido}, \
        @Estado = '{payload.Estado}', \
       """
    print (query)
    with engine.begin() as conn:
          conn.execution_options(autocommit = True)
          roles_df = pd.read_sql(query, conn)
    dumpp = ResponseModel(id_resultado=1,respuesta="El histórico del checklist se guardó de manera correcta")
    dict = dumpp.model_dump()
    return JSONResponse(status_code=200, content=dict)

@app.post(
        path="/api/checklist",
        name='Insertar checklist',
        tags=['Checklist'],
        description='Método para insertar el checklist',
        response_model=Checklist
)
def savechecklist(payload: Checklist):
    query = f"""
    EXEC InsertChecklist
        @lectura_codigos = {payload.lectura_codigos}, \
        @lectura_codigos_observacion = N'{payload.lectura_codigos_observacion}' , \
        @lectura_codigos_foto = N'{payload.lectura_codigos_foto}' , \
        @servofreno = {payload.servofreno}, \
        @servofreno_observacion = N'{payload.servofreno_observacion}' , \
        @servofreno_foto = N'{payload.servofreno_foto}' , \
        @pedal_freno = {payload.pedal_freno}, \
        @pedal_freno_observacion = N'{payload.pedal_freno_observacion}' , \
        @pedal_freno_foto = N'{payload.pedal_freno_foto}' , \
        @pedal_estacionamiento = {payload.pedal_estacionamiento}, \
        @pedal_estacionamiento_observacion = N'{payload.pedal_estacionamiento_observacion}' , \
        @pedal_estacionamiento_foto = N'{payload.pedal_estacionamiento_foto}' , \
        @cinturon_seguridad = {payload.cinturon_seguridad}, \
        @cinturon_seguridad_observacion = N'{payload.cinturon_seguridad_observacion}' , \
        @cinturon_seguridad_foto = N'{payload.cinturon_seguridad_foto}' , \
        @cuadro_instrumentos = {payload.cuadro_instrumentos}, \
        @cuadro_instrumentos_observacion = N'{payload.cuadro_instrumentos_observacion}' ,\
        @cuadro_instrumentos_foto = N'{payload.cuadro_instrumentos_foto}' ,\
        @aire_acondicionado = {payload.aire_acondicionado},\
        @aire_acondicionado_observacion = N'{payload.aire_acondicionado_observacion}' ,\
        @aire_acondicionado_foto = N'{payload.aire_acondicionado_foto}' ,\
        @bocina_claxon = {payload.bocina_claxon},\
        @bocina_claxon_observacion = N'{payload.bocina_claxon_observacion}' ,\
        @bocina_claxon_foto = N'{payload.bocina_claxon_foto}' ,\
        @iluminacion_interior = {payload.iluminacion_interior},\
        @iluminacion_interior_observacion = N'{payload.iluminacion_interior_observacion}' ,\
        @iluminacion_interior_foto = N'{payload.iluminacion_interior_foto}' ,\
        @iluminacion_externa = {payload.iluminacion_externa},\
        @iluminacion_externa_observacion = N'{payload.iluminacion_externa_observacion}' ,\
        @iluminacion_externa_foto = N'{payload.iluminacion_externa_foto}' ,\
        @limpiaparabrisas = {payload.limpiaparabrisas}, \
        @limpiaparabrisas_observacion = N'{payload.limpiaparabrisas_observacion}' ,\
        @limpiaparabrisas_foto = N'{payload.limpiaparabrisas_foto}' , \
        @limpia_medallon = {payload.limpia_medallon}, \
        @limpia_medallon_observacion = N'{payload.limpia_medallon_observacion}' , \
        @limpia_medallon_foto = N'{payload.limpia_medallon_foto}' , \
        @neumaticos_friccion = {payload.neumaticos_friccion}, \
        @neumaticos_friccion_observacion = N'{payload.neumaticos_friccion_observacion}' ,  \
        @neumaticos_friccion_foto = N'{payload.neumaticos_friccion_foto}' , \
        @otros_vehiculo_en_piso = {payload.otros_vehiculo_en_piso}, \
        @otros_vehiculo_en_piso_observacion = N'{payload.otros_vehiculo_en_piso_observacion}' , \
        @otros_vehiculo_en_piso_foto = N'{payload.otros_vehiculo_en_piso_foto}' , \
        @estado_fugas_aceite = {payload.estado_fugas_aceite}, \
        @estado_fugas_aceite_observacion = N'{payload.estado_fugas_aceite_observacion}' , \
        @estado_fugas_aceite_foto = N'{payload.estado_fugas_aceite_foto}' , \
        @estado_nivel_calidad_lubricante_transmision = {payload.estado_nivel_calidad_lubricante_transmision}, \
        @estado_nivel_calidad_lubricante_transmision_observacion = N'{payload.estado_nivel_calidad_lubricante_transmision_observacion}' , \
        @estado_nivel_calidad_lubricante_transmision_foto = N'{payload.estado_nivel_calidad_lubricante_transmision_foto}' , \
        @estado_nivel_calidad_lubricante_diferencial = {payload.estado_nivel_calidad_lubricante_diferencial}, \
        @estado_nivel_calidad_lubricante_diferencial_observacion = N'{payload.estado_nivel_calidad_lubricante_diferencial_observacion}' , \
        @estado_nivel_calidad_lubricante_diferencial_foto = N'{payload.estado_nivel_calidad_lubricante_diferencial_foto}' , \
        @cubrepolvos_flechas = {payload.cubrepolvos_flechas}, \
        @cubrepolvos_flechas_observacion = N'{payload.cubrepolvos_flechas_observacion}' , \
        @cubrepolvos_flechas_foto = N'{payload.cubrepolvos_flechas_foto}' , \
        @componentes_direccion = {payload.componentes_direccion}, \
        @componentes_direccion_observacion = N'{payload.componentes_direccion_observacion}' , \
        @componentes_direccion_foto = N'{payload.componentes_direccion_foto}' , \
        @componentes_suspesion = {payload.componentes_suspesion}, \
        @componentes_suspesion_observacion = N'{payload.componentes_suspesion_observacion}' , \
        @componentes_suspesion_foto = N'{payload.componentes_suspesion_foto}' , \
        @sistema_escape_completo = {payload.sistema_escape_completo}, \
        @sistema_escape_completo_observacion = N'{payload.sistema_escape_completo_observacion}' , \
        @sistema_escape_completo_foto = N'{payload.sistema_escape_completo_foto}' , \
        @sistema_alimentacion_combustible = {payload.sistema_alimentacion_combustible}, \
        @sistema_alimentacion_combustible_observacion = N'{payload.sistema_alimentacion_combustible_observacion}' , \
        @sistema_alimentacion_combustible_foto = N'{payload.sistema_alimentacion_combustible_foto}' , \
        @filtro_combustible = {payload.filtro_combustible}, \
        @filtro_combustible_observacion = N'{payload.filtro_combustible_observacion}' , \
        @filtro_combustible_foto = N'{payload.filtro_combustible_foto}' , \
        @control_fugas_direccion_hidraulica = {payload.control_fugas_direccion_hidraulica}, \
        @control_fugas_direccion_hidraulica_observacion = N'{payload.control_fugas_direccion_hidraulica_observacion}' , \
        @control_fugas_direccion_hidraulica_foto = N'{payload.control_fugas_direccion_hidraulica_foto}' , \
        @otros_altura_total = {payload.otros_altura_total}, \
        @otros_altura_total_observacion = N'{payload.otros_altura_total_observacion}' , \
        @otros_altura_total_foto = N'{payload.otros_altura_total_foto}' , \
        @rodamiento_mazas_rueda = {payload.rodamiento_mazas_rueda}, \
        @rodamiento_mazas_rueda_observacion = N'{payload.rodamiento_mazas_rueda_observacion}' , \
        @rodamiento_mazas_rueda_foto = N'{payload.rodamiento_mazas_rueda_foto}' , \
        @holgura_partes_suspension_rueda = {payload.holgura_partes_suspension_rueda}, \
        @holgura_partes_suspension_rueda_observacion = N'{payload.holgura_partes_suspension_rueda_observacion}' , \
        @holgura_partes_suspension_rueda_foto = N'{payload.holgura_partes_suspension_rueda_foto}' , \
        @control_neumaticos_desgaste_presion = {payload.control_neumaticos_desgaste_presion}, \
        @control_neumaticos_desgaste_presion_observacion = N'{payload.control_neumaticos_desgaste_presion_observacion}' , \
        @control_neumaticos_desgaste_presion_foto = N'{payload.control_neumaticos_desgaste_presion_foto}' , \
        @profundidad = {payload.profundidad}, \
        @profundidad_observacion = N'{payload.profundidad_observacion}' , \
        @profundidad_foto = N'{payload.profundidad_foto}' , \
        @presion = {payload.presion}, \
        @presion_observacion = N'{payload.presion_observacion}' , \
        @presion_foto = N'{payload.presion_foto}' , \
        @otros_altura_media = {payload.otros_altura_media}, \
        @otros_altura_media_observacion = N'{payload.otros_altura_media_observacion}' , \
        @otros_altura_media_foto = N'{payload.otros_altura_media_foto}' , \
        @nivel_calidad_aceite_motor = {payload.nivel_calidad_aceite_motor}, \
        @nivel_calidad_aceite_motor_observacion = N'{payload.nivel_calidad_aceite_motor_observacion}' , \
        @nivel_calidad_aceite_motor_foto = N'{payload.nivel_calidad_aceite_motor_foto}' , \
        @filtro_aire = {payload.filtro_aire}, \
        @filtro_aire_observacion = N'{payload.filtro_aire_observacion}' , \
        @filtro_aire_foto = N'{payload.filtro_aire_foto}' , \
        @filtro_polen = {payload.filtro_polen}, \
        @filtro_polen_observacion = N'{payload.filtro_polen_observacion}' , \
        @filtro_polen_foto = N'{payload.filtro_polen_foto}' , \
        @filtro_pcv = {payload.filtro_pcv}, \
        @filtro_pcv_observacion = N'{payload.filtro_pcv_observacion}' , \
        @filtro_pcv_foto = N'{payload.filtro_pcv_foto}' , \
        @valvula_pcv = {payload.valvula_pcv}, \
        @valvula_pcv_observacion = N'{payload.valvula_pcv_observacion}' , \
        @valvula_pcv_foto = N'{payload.valvula_pcv_foto}' , \
        @bujias_encendido = {payload.bujias_encendido}, \
        @bujias_encendido_observacion = N'{payload.bujias_encendido_observacion}' , \
        @bujias_encendido_foto = N'{payload.bujias_encendido_foto}' , \
        @cables_bujias_bobinas_ignicion = {payload.cables_bujias_bobinas_ignicion}, \
        @cables_bujias_bobinas_ignicion_observacion = N'{payload.cables_bujias_bobinas_ignicion_observacion}' , \
        @cables_bujias_bobinas_ignicion_foto = N'{payload.cables_bujias_bobinas_ignicion_foto}' , \
        @nivel_anticongenlante = {payload.nivel_anticongenlante}, \
        @nivel_anticongenlante_observacion = N'{payload.nivel_anticongenlante_observacion}' , \
        @nivel_anticongenlante_foto = N'{payload.nivel_anticongenlante_foto}' , \
        @tapon_radiador = {payload.tapon_radiador}, \
        @tapon_radiador_observacion = N'{payload.tapon_radiador_observacion}' , \
        @tapon_radiador_foto = N'{payload.tapon_radiador_foto}' , \
        @mangueras_sistema = {payload.mangueras_sistema}, \
        @mangueras_sistema_observacion = N'{payload.mangueras_sistema_observacion}' , \
        @mangueras_sistema_foto = N'{payload.mangueras_sistema_foto}' , \
        @desempeño_ventilador = {payload.desempeño_ventilador}, \
        @desempeño_ventilador_observacion = N'{payload.desempeño_ventilador_observacion}' , \
        @desempeño_ventilador_foto = N'{payload.desempeño_ventilador_foto}' , \
        @calidad_liquido_limpiaparabrisas = {payload.calidad_liquido_limpiaparabrisas}, \
        @calidad_liquido_limpiaparabrisas_observacion = N'{payload.calidad_liquido_limpiaparabrisas_observacion}' , \
        @calidad_liquido_limpiaparabrisas_foto = N'{payload.calidad_liquido_limpiaparabrisas_foto}' , \
        @calidad_aceite_direccion_hidraulica = {payload.calidad_aceite_direccion_hidraulica}, \
        @calidad_aceite_direccion_hidraulica_observacion = N'{payload.calidad_aceite_direccion_hidraulica_observacion}' , \
        @calidad_aceite_direccion_hidraulica_foto = N'{payload.calidad_aceite_direccion_hidraulica_foto}' , \
        @calidad_aceite_transmision_bayoneta = {payload.calidad_aceite_transmision_bayoneta}, \
        @calidad_aceite_transmision_bayoneta_observacion = N'{payload.calidad_aceite_transmision_bayoneta_observacion}' , \
        @calidad_aceite_transmision_bayoneta_foto = N'{payload.calidad_aceite_transmision_bayoneta_foto}' , \
        @liquido_bateria_condiciones = {payload.liquido_bateria_condiciones}, \
        @liquido_bateria_condiciones_observacion = N'{payload.liquido_bateria_condiciones_observacion}' , \
        @liquido_bateria_condiciones_foto = N'{payload.liquido_bateria_condiciones_foto}' ,    \
        @bandas_poly_v = {payload.bandas_poly_v}, \
        @bandas_poly_v_observacion = N'{payload.bandas_poly_v_observacion}' , \
        @bandas_poly_v_foto = N'{payload.bandas_poly_v_foto}' , \
        @poleas_banda = {payload.poleas_banda}, \
        @poleas_banda_observacion = N'{payload.poleas_banda_observacion}' , \
        @poleas_banda_foto = N'{payload.poleas_banda_foto}' , \
        @banda_tiempo = {payload.banda_tiempo}, \
        @banda_tiempo_observacion = N'{payload.banda_tiempo_observacion}' , \
        @banda_tiempo_foto = N'{payload.banda_tiempo_foto}' , \
        @otros_habitaculo_motor = {payload.otros_habitaculo_motor}, \
        @otros_habitaculo_motor_observacion = N'{payload.otros_habitaculo_motor_observacion}' , \
        @otros_habitaculo_motor_foto = N'{payload.otros_habitaculo_motor_foto}' , \
        @reset_intervalo_servicio = {payload.reset_intervalo_servicio}, \
        @reset_intervalo_servicio_observacion = N'{payload.reset_intervalo_servicio_observacion}' , \
        @reset_intervalo_servicio_foto = N'{payload.reset_intervalo_servicio_foto}' , \
        @ajuste_tornillos_neumaticos_torquimetro = {payload.ajuste_tornillos_neumaticos_torquimetro}, \
        @ajuste_tornillos_neumaticos_torquimetro_observacion = N'{payload.ajuste_tornillos_neumaticos_torquimetro_observacion}' , \
        @ajuste_tornillos_neumaticos_torquimetro_foto = N'{payload.ajuste_tornillos_neumaticos_torquimetro_foto}' , \
        @limpiar_libricar_puertas_cerraduras = {payload.limpiar_libricar_puertas_cerraduras}, \
        @limpiar_libricar_puertas_cerraduras_observacion = N'{payload.limpiar_libricar_puertas_cerraduras_observacion}' , \
        @limpiar_libricar_puertas_cerraduras_foto = N'{payload.limpiar_libricar_puertas_cerraduras_foto}' , \
        @completar_plan_mantenimiento = {payload.completar_plan_mantenimiento}, \
        @completar_plan_mantenimiento_observacion = N'{payload.completar_plan_mantenimiento_observacion}' , \
        @completar_plan_mantenimiento_foto = N'{payload.completar_plan_mantenimiento_foto}' , \
        @fecha = N'{payload.Fecha}' , \
        @Id_empleado = {payload.IdEmpleado} , \
        @Id_vehiculo = {payload.IdVehiculo} , \
        @id_ordendeservicio = {payload.Id_ordendeservicio} , \
        @NumeroSerie = '{payload.NumeroSerie}' """
    print (query)
    with engine.begin() as conn:
          conn.execution_options(autocommit = True)
          roles_df = pd.read_sql(query, conn)
    dumpp = ResponseModel(id_resultado=1,respuesta="El checklist se guardó de manera correcta")
    dict = dumpp.model_dump()
    return JSONResponse(status_code=200, content=dict)



@app.post(
        path="/api/servicio",
        name='Insertar servicio',
        tags=['Servicio'],
        description='Método para insertar el servicio',
        response_model=Checklist
)
def saveservicio(payload: Checklist):
    query = f"""
    EXEC InsertServicio
        @lectura_codigos = {payload.lectura_codigos}, \
        @lectura_codigos_observacion = N'{payload.lectura_codigos_observacion}' , \
        @lectura_codigos_foto = N'{payload.lectura_codigos_foto}' , \
        @servofreno = {payload.servofreno}, \
        @servofreno_observacion = N'{payload.servofreno_observacion}' , \
        @servofreno_foto = N'{payload.servofreno_foto}' , \
        @pedal_freno = {payload.pedal_freno}, \
        @pedal_freno_observacion = N'{payload.pedal_freno_observacion}' , \
        @pedal_freno_foto = N'{payload.pedal_freno_foto}' , \
        @pedal_estacionamiento = {payload.pedal_estacionamiento}, \
        @pedal_estacionamiento_observacion = N'{payload.pedal_estacionamiento_observacion}' , \
        @pedal_estacionamiento_foto = N'{payload.pedal_estacionamiento_foto}' , \
        @cinturon_seguridad = {payload.cinturon_seguridad}, \
        @cinturon_seguridad_observacion = N'{payload.cinturon_seguridad_observacion}' , \
        @cinturon_seguridad_foto = N'{payload.cinturon_seguridad_foto}' , \
        @cuadro_instrumentos = {payload.cuadro_instrumentos}, \
        @cuadro_instrumentos_observacion = N'{payload.cuadro_instrumentos_observacion}' ,\
        @cuadro_instrumentos_foto = N'{payload.cuadro_instrumentos_foto}' ,\
        @aire_acondicionado = {payload.aire_acondicionado},\
        @aire_acondicionado_observacion = N'{payload.aire_acondicionado_observacion}' ,\
        @aire_acondicionado_foto = N'{payload.aire_acondicionado_foto}' ,\
        @bocina_claxon = {payload.bocina_claxon},\
        @bocina_claxon_observacion = N'{payload.bocina_claxon_observacion}' ,\
        @bocina_claxon_foto = N'{payload.bocina_claxon_foto}' ,\
        @iluminacion_interior = {payload.iluminacion_interior},\
        @iluminacion_interior_observacion = N'{payload.iluminacion_interior_observacion}' ,\
        @iluminacion_interior_foto = N'{payload.iluminacion_interior_foto}' ,\
        @iluminacion_externa = {payload.iluminacion_externa},\
        @iluminacion_externa_observacion = N'{payload.iluminacion_externa_observacion}' ,\
        @iluminacion_externa_foto = N'{payload.iluminacion_externa_foto}' ,\
        @limpiaparabrisas = {payload.limpiaparabrisas}, \
        @limpiaparabrisas_observacion = N'{payload.limpiaparabrisas_observacion}' ,\
        @limpiaparabrisas_foto = N'{payload.limpiaparabrisas_foto}' , \
        @limpia_medallon = {payload.limpia_medallon}, \
        @limpia_medallon_observacion = N'{payload.limpia_medallon_observacion}' , \
        @limpia_medallon_foto = N'{payload.limpia_medallon_foto}' , \
        @neumaticos_friccion = {payload.neumaticos_friccion}, \
        @neumaticos_friccion_observacion = N'{payload.neumaticos_friccion_observacion}' ,  \
        @neumaticos_friccion_foto = N'{payload.neumaticos_friccion_foto}' , \
        @otros_vehiculo_en_piso = {payload.otros_vehiculo_en_piso}, \
        @otros_vehiculo_en_piso_observacion = N'{payload.otros_vehiculo_en_piso_observacion}' , \
        @otros_vehiculo_en_piso_foto = N'{payload.otros_vehiculo_en_piso_foto}' , \
        @estado_fugas_aceite = {payload.estado_fugas_aceite}, \
        @estado_fugas_aceite_observacion = N'{payload.estado_fugas_aceite_observacion}' , \
        @estado_fugas_aceite_foto = N'{payload.estado_fugas_aceite_foto}' , \
        @estado_nivel_calidad_lubricante_transmision = {payload.estado_nivel_calidad_lubricante_transmision}, \
        @estado_nivel_calidad_lubricante_transmision_observacion = N'{payload.estado_nivel_calidad_lubricante_transmision_observacion}' , \
        @estado_nivel_calidad_lubricante_transmision_foto = N'{payload.estado_nivel_calidad_lubricante_transmision_foto}' , \
        @estado_nivel_calidad_lubricante_diferencial = {payload.estado_nivel_calidad_lubricante_diferencial}, \
        @estado_nivel_calidad_lubricante_diferencial_observacion = N'{payload.estado_nivel_calidad_lubricante_diferencial_observacion}' , \
        @estado_nivel_calidad_lubricante_diferencial_foto = N'{payload.estado_nivel_calidad_lubricante_diferencial_foto}' , \
        @cubrepolvos_flechas = {payload.cubrepolvos_flechas}, \
        @cubrepolvos_flechas_observacion = N'{payload.cubrepolvos_flechas_observacion}' , \
        @cubrepolvos_flechas_foto = N'{payload.cubrepolvos_flechas_foto}' , \
        @componentes_direccion = {payload.componentes_direccion}, \
        @componentes_direccion_observacion = N'{payload.componentes_direccion_observacion}' , \
        @componentes_direccion_foto = N'{payload.componentes_direccion_foto}' , \
        @componentes_suspesion = {payload.componentes_suspesion}, \
        @componentes_suspesion_observacion = N'{payload.componentes_suspesion_observacion}' , \
        @componentes_suspesion_foto = N'{payload.componentes_suspesion_foto}' , \
        @sistema_escape_completo = {payload.sistema_escape_completo}, \
        @sistema_escape_completo_observacion = N'{payload.sistema_escape_completo_observacion}' , \
        @sistema_escape_completo_foto = N'{payload.sistema_escape_completo_foto}' , \
        @sistema_alimentacion_combustible = {payload.sistema_alimentacion_combustible}, \
        @sistema_alimentacion_combustible_observacion = N'{payload.sistema_alimentacion_combustible_observacion}' , \
        @sistema_alimentacion_combustible_foto = N'{payload.sistema_alimentacion_combustible_foto}' , \
        @filtro_combustible = {payload.filtro_combustible}, \
        @filtro_combustible_observacion = N'{payload.filtro_combustible_observacion}' , \
        @filtro_combustible_foto = N'{payload.filtro_combustible_foto}' , \
        @control_fugas_direccion_hidraulica = {payload.control_fugas_direccion_hidraulica}, \
        @control_fugas_direccion_hidraulica_observacion = N'{payload.control_fugas_direccion_hidraulica_observacion}' , \
        @control_fugas_direccion_hidraulica_foto = N'{payload.control_fugas_direccion_hidraulica_foto}' , \
        @otros_altura_total = {payload.otros_altura_total}, \
        @otros_altura_total_observacion = N'{payload.otros_altura_total_observacion}' , \
        @otros_altura_total_foto = N'{payload.otros_altura_total_foto}' , \
        @rodamiento_mazas_rueda = {payload.rodamiento_mazas_rueda}, \
        @rodamiento_mazas_rueda_observacion = N'{payload.rodamiento_mazas_rueda_observacion}' , \
        @rodamiento_mazas_rueda_foto = N'{payload.rodamiento_mazas_rueda_foto}' , \
        @holgura_partes_suspension_rueda = {payload.holgura_partes_suspension_rueda}, \
        @holgura_partes_suspension_rueda_observacion = N'{payload.holgura_partes_suspension_rueda_observacion}' , \
        @holgura_partes_suspension_rueda_foto = N'{payload.holgura_partes_suspension_rueda_foto}' , \
        @control_neumaticos_desgaste_presion = {payload.control_neumaticos_desgaste_presion}, \
        @control_neumaticos_desgaste_presion_observacion = N'{payload.control_neumaticos_desgaste_presion_observacion}' , \
        @control_neumaticos_desgaste_presion_foto = N'{payload.control_neumaticos_desgaste_presion_foto}' , \
        @profundidad = {payload.profundidad}, \
        @profundidad_observacion = N'{payload.profundidad_observacion}' , \
        @profundidad_foto = N'{payload.profundidad_foto}' , \
        @presion = {payload.presion}, \
        @presion_observacion = N'{payload.presion_observacion}' , \
        @presion_foto = N'{payload.presion_foto}' , \
        @otros_altura_media = {payload.otros_altura_media}, \
        @otros_altura_media_observacion = N'{payload.otros_altura_media_observacion}' , \
        @otros_altura_media_foto = N'{payload.otros_altura_media_foto}' , \
        @nivel_calidad_aceite_motor = {payload.nivel_calidad_aceite_motor}, \
        @nivel_calidad_aceite_motor_observacion = N'{payload.nivel_calidad_aceite_motor_observacion}' , \
        @nivel_calidad_aceite_motor_foto = N'{payload.nivel_calidad_aceite_motor_foto}' , \
        @filtro_aire = {payload.filtro_aire}, \
        @filtro_aire_observacion = N'{payload.filtro_aire_observacion}' , \
        @filtro_aire_foto = N'{payload.filtro_aire_foto}' , \
        @filtro_polen = {payload.filtro_polen}, \
        @filtro_polen_observacion = N'{payload.filtro_polen_observacion}' , \
        @filtro_polen_foto = N'{payload.filtro_polen_foto}' , \
        @filtro_pcv = {payload.filtro_pcv}, \
        @filtro_pcv_observacion = N'{payload.filtro_pcv_observacion}' , \
        @filtro_pcv_foto = N'{payload.filtro_pcv_foto}' , \
        @valvula_pcv = {payload.valvula_pcv}, \
        @valvula_pcv_observacion = N'{payload.valvula_pcv_observacion}' , \
        @valvula_pcv_foto = N'{payload.valvula_pcv_foto}' , \
        @bujias_encendido = {payload.bujias_encendido}, \
        @bujias_encendido_observacion = N'{payload.bujias_encendido_observacion}' , \
        @bujias_encendido_foto = N'{payload.bujias_encendido_foto}' , \
        @cables_bujias_bobinas_ignicion = {payload.cables_bujias_bobinas_ignicion}, \
        @cables_bujias_bobinas_ignicion_observacion = N'{payload.cables_bujias_bobinas_ignicion_observacion}' , \
        @cables_bujias_bobinas_ignicion_foto = N'{payload.cables_bujias_bobinas_ignicion_foto}' , \
        @nivel_anticongenlante = {payload.nivel_anticongenlante}, \
        @nivel_anticongenlante_observacion = N'{payload.nivel_anticongenlante_observacion}' , \
        @nivel_anticongenlante_foto = N'{payload.nivel_anticongenlante_foto}' , \
        @tapon_radiador = {payload.tapon_radiador}, \
        @tapon_radiador_observacion = N'{payload.tapon_radiador_observacion}' , \
        @tapon_radiador_foto = N'{payload.tapon_radiador_foto}' , \
        @mangueras_sistema = {payload.mangueras_sistema}, \
        @mangueras_sistema_observacion = N'{payload.mangueras_sistema_observacion}' , \
        @mangueras_sistema_foto = N'{payload.mangueras_sistema_foto}' , \
        @desempeño_ventilador = {payload.desempeño_ventilador}, \
        @desempeño_ventilador_observacion = N'{payload.desempeño_ventilador_observacion}' , \
        @desempeño_ventilador_foto = N'{payload.desempeño_ventilador_foto}' , \
        @calidad_liquido_limpiaparabrisas = {payload.calidad_liquido_limpiaparabrisas}, \
        @calidad_liquido_limpiaparabrisas_observacion = N'{payload.calidad_liquido_limpiaparabrisas_observacion}' , \
        @calidad_liquido_limpiaparabrisas_foto = N'{payload.calidad_liquido_limpiaparabrisas_foto}' , \
        @calidad_aceite_direccion_hidraulica = {payload.calidad_aceite_direccion_hidraulica}, \
        @calidad_aceite_direccion_hidraulica_observacion = N'{payload.calidad_aceite_direccion_hidraulica_observacion}' , \
        @calidad_aceite_direccion_hidraulica_foto = N'{payload.calidad_aceite_direccion_hidraulica_foto}' , \
        @calidad_aceite_transmision_bayoneta = {payload.calidad_aceite_transmision_bayoneta}, \
        @calidad_aceite_transmision_bayoneta_observacion = N'{payload.calidad_aceite_transmision_bayoneta_observacion}' , \
        @calidad_aceite_transmision_bayoneta_foto = N'{payload.calidad_aceite_transmision_bayoneta_foto}' , \
        @liquido_bateria_condiciones = {payload.liquido_bateria_condiciones}, \
        @liquido_bateria_condiciones_observacion = N'{payload.liquido_bateria_condiciones_observacion}' , \
        @liquido_bateria_condiciones_foto = N'{payload.liquido_bateria_condiciones_foto}' ,    \
        @bandas_poly_v = {payload.bandas_poly_v}, \
        @bandas_poly_v_observacion = N'{payload.bandas_poly_v_observacion}' , \
        @bandas_poly_v_foto = N'{payload.bandas_poly_v_foto}' , \
        @poleas_banda = {payload.poleas_banda}, \
        @poleas_banda_observacion = N'{payload.poleas_banda_observacion}' , \
        @poleas_banda_foto = N'{payload.poleas_banda_foto}' , \
        @banda_tiempo = {payload.banda_tiempo}, \
        @banda_tiempo_observacion = N'{payload.banda_tiempo_observacion}' , \
        @banda_tiempo_foto = N'{payload.banda_tiempo_foto}' , \
        @otros_habitaculo_motor = {payload.otros_habitaculo_motor}, \
        @otros_habitaculo_motor_observacion = N'{payload.otros_habitaculo_motor_observacion}' , \
        @otros_habitaculo_motor_foto = N'{payload.otros_habitaculo_motor_foto}' , \
        @reset_intervalo_servicio = {payload.reset_intervalo_servicio}, \
        @reset_intervalo_servicio_observacion = N'{payload.reset_intervalo_servicio_observacion}' , \
        @reset_intervalo_servicio_foto = N'{payload.reset_intervalo_servicio_foto}' , \
        @ajuste_tornillos_neumaticos_torquimetro = {payload.ajuste_tornillos_neumaticos_torquimetro}, \
        @ajuste_tornillos_neumaticos_torquimetro_observacion = N'{payload.ajuste_tornillos_neumaticos_torquimetro_observacion}' , \
        @ajuste_tornillos_neumaticos_torquimetro_foto = N'{payload.ajuste_tornillos_neumaticos_torquimetro_foto}' , \
        @limpiar_libricar_puertas_cerraduras = {payload.limpiar_libricar_puertas_cerraduras}, \
        @limpiar_libricar_puertas_cerraduras_observacion = N'{payload.limpiar_libricar_puertas_cerraduras_observacion}' , \
        @limpiar_libricar_puertas_cerraduras_foto = N'{payload.limpiar_libricar_puertas_cerraduras_foto}' , \
        @completar_plan_mantenimiento = {payload.completar_plan_mantenimiento}, \
        @completar_plan_mantenimiento_observacion = N'{payload.completar_plan_mantenimiento_observacion}' , \
        @completar_plan_mantenimiento_foto = N'{payload.completar_plan_mantenimiento_foto}' , \
        @fecha = N'{payload.Fecha}' , \
        @Id_empleado = {payload.IdEmpleado} , \
        @Id_vehiculo = {payload.IdVehiculo} , \
        @id_ordendeservicio = {payload.Id_ordendeservicio} , \
        @id_checklist = {payload.id_checklist} , \
        @NumeroSerie = '{payload.NumeroSerie}' """
    print (query)
    with engine.begin() as conn:
          conn.execution_options(autocommit = True)
          roles_df = pd.read_sql(query, conn)
    dumpp = ResponseModel(id_resultado=1,respuesta="El servicio se guardó de manera correcta")
    dict = dumpp.model_dump()
    return JSONResponse(status_code=200, content=dict)


@app.get(
        path="/api/obtenerservicio",
        name='Obtener servicio',
        tags=['Servicio'],
        description='Método para obtener la informacion de 1 checklist',
        response_model=Checklist
)
def getservicio(Idchecklist: int):
    query = f"exec [dbo].[sp_get_all_servicio] @IdCheckList = {Idchecklist}"
    roles_df = pd.read_sql(query, engine)
    resultado = roles_df.to_dict(orient="records")
    print(resultado)
    return JSONResponse(status_code=200,content=resultado[0])


@app.get(
        path="/api/obtenerflotillaporid",
        name='Obtener flotilla por id',
        tags=['Flotillas'],
        description='Método para obtener la informacion de 1 flotilla',
        response_model=Flotillas
)
def getflotillaevidencia(IdFlotilla: int):
    query = f"exec [dbo].[ObtenerAllFlotillasPorID] @IdFlotilla = {IdFlotilla}"
    roles_df = pd.read_sql(query, engine)
    resultado = roles_df.to_dict(orient="records")
    print(resultado)
    return JSONResponse(status_code=200,content=resultado[0])

@app.get(
        path="/api/obtenerallflotillas",
        name='Obtener todas las flotillas',
        tags=['Flotillas'],
        description='Método para obtener la informacion de todos las flotillas',
        response_model=Flotillas
)
def getsallflotillas():
    query = f"exec [dbo].[ObtenerAllFlotillas]"
    roles_df = pd.read_sql(query, engine)
    resultado = roles_df.to_dict(orient="records")
    return JSONResponse(status_code=200,content=resultado)


@app.get(
        path="/api/obtenerhistoricos",
        name='Obtener históricos',
        tags=['Historico'],
        description='Método para obtener la informacion de todos los históricos',
        response_model=CheckListHistorico
)
def getshistoricocheck():
    query = f"exec [dbo].[ObtenerHistoricoCheck]"
    roles_df = pd.read_sql(query, engine)
    resultado = roles_df.to_dict(orient="records")
    return JSONResponse(status_code=200,content=resultado)

@app.get(
        path="/api/obtenerhistoricosservicios",
        name='Obtener históricos servicios',
        tags=['Historico'],
        description='Método para obtener la informacion de todos los históricos de servicios',
        response_model=CheckListHistorico
)
def getshistoricoservicio():
    query = f"exec [dbo].[ObtenerHistoricoServicio]"
    roles_df = pd.read_sql(query, engine)
    resultado = roles_df.to_dict(orient="records")
    return JSONResponse(status_code=200,content=resultado)

@app.get(
        path="/api/obtenerflotillas",
        name='Obtener flotillas',
        tags=['Flotillas'],
        description='Método para obtener la informacion de todas las flotillas',
        response_model=Flotillas
)
def getsflotillas():
    query = f"exec [dbo].[ObtenerFlotillas]"
    roles_df = pd.read_sql(query, engine)
    resultado = roles_df.to_dict(orient="records")
    return JSONResponse(status_code=200,content=resultado)

@app.get(
        path="/api/obtenerflotilla",
        name='Obtener flotilla',
        tags=['Flotillas'],
        description='Método para obtener la informacion de 1 flotilla',
        response_model=Flotillas
)
def getflotilla(IdFlotilla: int):
    query = f"exec [dbo].[sp_get_all_flotillas] @IdFlotilla = {IdFlotilla}"
    roles_df = pd.read_sql(query, engine)
    resultado = roles_df.to_dict(orient="records")
    print(resultado)
    return JSONResponse(status_code=200,content=resultado[0])

@app.post(
    path="/api/flotilla",
    name="Guarda Flotilla",
    tags=["Flotillas"],
    description="Método para guardar la información de las flotillas",
    response_model=Flotillas,
)
def guardarFlotilla(payload: Flotillas):
    try:
        # Crear el diccionario de parámetros desde el payload
        parametros = payload.dict()

        # Definir la consulta SQL para ejecutar el procedimiento almacenado
        query = text("""
            EXEC dbo.Insertflotillas 
                @NamesFlotillas = :NamesFlotillas,
                @Encargado = :Encargado
        """)

        # Ejecutar la consulta SQL
        with engine.begin() as conn:
            conn.execute(query, parametros)

        # Respuesta de éxito
        return JSONResponse(
            status_code=200,
            content={
                "id_resultado": 1,
                "respuesta": "La flotilla se guardó de manera correcta",
                "detalles": parametros,
            },
        )

    except Exception as e:
        # Respuesta de error
        raise HTTPException(
            status_code=500,
            detail=f"Error al guardar la flotilla: {str(e)}",
        )


@app.get(
        path="/api/obtenerservicios",
        name='Obtener servicios',
        tags=['Servicio'],
        description='Método para obtener la informacion todos los servicios',
        response_model=Checklist
)
def getservicios():
    query = f"exec [dbo].[ObtenerServicios]"
    roles_df = pd.read_sql(query, engine)
    resultado = roles_df.to_dict(orient="records")
    return JSONResponse(status_code=200,content=resultado)

@app.get(
        path="/api/obtenerchecklist",
        name='Obtener checklist',
        tags=['Checklist'],
        description='Método para obtener la informacion de 1 checklist',
        response_model=Checklist
)
def getempleados(Idchecklist: int):
    query = f"exec [dbo].[sp_get_all_checklist] @IdCheckList = {Idchecklist}"
    roles_df = pd.read_sql(query, engine)
    resultado = roles_df.to_dict(orient="records")
    print(resultado)
    return JSONResponse(status_code=200,content=resultado[0])


@app.get(
        path="/api/obtenerchecklists",
        name='Obtener checklists',
        tags=['Checklist'],
        description='Método para obtener la informacion todos los checklist',
        response_model=Checklist
)
def getchecklists():
    query = f"exec [dbo].[ObtenerCheckLists]"
    roles_df = pd.read_sql(query, engine)
    resultado = roles_df.to_dict(orient="records")
    return JSONResponse(status_code=200,content=resultado)
@app.get(
    path="/api/obteneridOrden",
    name="Obtener ID de Orden de Servicio",
    tags=["Orden"],
    description="Obtiene el siguiente ID secuencial para la orden de servicio."
)
def obtener_id_orden():
    query = f"exec [dbo].ObtenerUltimoIdOrdenServicio"
    roles_df = pd.read_sql(query, engine)
    resultado = roles_df.to_dict(orient="records")
    print(resultado)
    return JSONResponse(status_code=200,content=resultado[0])

@app.get(
    path="/api/obteneridCheck",
    name='Obtener IDs del Checklist',
    tags=['Checklist'],
    description='Obtiene todos los IDs de los checklists existentes',
)
def obtener_id_check():
    query = "exec [dbo].ObtenerIdCheckList"
    roles_df = pd.read_sql(query, engine)
    resultado = roles_df.to_dict(orient="records")
    return JSONResponse(status_code=200, content=resultado)


@app.get(
    path="/api/obtenerchecklisthtml",
    name='Obtener checklist HTML',
    tags=['Checklist'],
    description='Método para obtener el HTML del checklist',
)
def obtener_checklist_html(Idchecklist: int):
    
    query = f"exec [dbo].[sp_get_all_checklist] @IdCheckList = {Idchecklist}"
    with engine.begin() as conn:
      conn.execution_options(autocommit=True)
    roles_df = pd.read_sql(query, conn)

        # Asegúrate de que hay al menos dos filas en el DataFrame
        
            # Accede al valor en la segunda fila de la columna 'Nombre'
            

            # Construye el HTML con el valor obtenido
    id = roles_df['id'].iloc[0]
    fecha = ['fecha'].iloc[0]
    idEmpleado = ['Id_empleado'].iloc[0]
    idVehiculo = ['Id_vehiculo'].iloc[0]
    idOrden= ['id_ordendeservicio'].iloc[0]
    lecturaCodigo = ['lectura_codigos'].iloc[0]
    lecturaCodigoObservacion =['lectura_codigos_observacion'].iloc[0]
    servofreno = ['servofreno'].iloc[0]
    servofrenoObservacion =['servofreno_observacion'].iloc[0]
    pedalFreno =['pedal_freno'].iloc[0]
    pedalFrenoObservacion =['pedal_freno_observacion'].iloc[0]
    pedalEstacionamiento =['pedal_estacionamiento'].iloc[0]
    pedalEstacionamientoObservacion=['pedal_estacionamiento_observacion'].iloc[0]
    cinturonSeguridad =['cinturon_seguridad'].iloc[0]
    cinturonSeguridadObservacion=['cinturon_seguridad_observacion'].iloc[0]
    cuadroInstrumentos=['cuadro_instrumentos'].iloc[0]
    cuadroInstrumentosObservacion=['cuadro_instrumentos_observacion'].iloc[0]
    aireAcondicionado =['aire_acondicionado'].iloc[0]
    aireAcondicionadoObservacion=['aire_acondicionado_observacion'].iloc[0]
    bocinaClaxon=['bocina_claxon'].iloc[0]
    bocinaClaxonObservacion=['bocina_claxon_observacion'].iloc[0]
    iluminacionInterior=['iluminacion_interior'].iloc[0]
    iluminacionInteriorObservacion=['iluminacion_interior_observacion'].iloc[0]
    iluminacionExterna=['iluminacion_externa'].iloc[0]
    iluminacionExternaObservacion=['iluminacion_externa_observacion'].iloc[0]
    limpiaparabrisas=['limpiaparabrisas'].iloc[0]
    limpiaparabrisasObservacion=['limpiaparabrisas_observacion'].iloc[0]
    limpiaMedallon=['limpia_medallon'].iloc[0]
    limpiaMedallonObservacion=['limpia_medallon_observacion'].iloc[0]
    neumaticosFriccion=['neumaticos_friccion'].iloc[0]
    neumaticosFriccionObservacion=['neumaticos_friccion_observacion'].iloc[0]
    otroVehiculosPiso=['otros_vehiculo_en_piso'].iloc[0]
    otroVehiculosPisoObservacion=['otros_vehiculo_en_piso_observacion'].iloc[0]
    estadoFugasAceite=['estado_fugas_aceite'].iloc[0]
    estadoFugasAceiteObservacion=['estado_fugas_aceite_observacion'].iloc[0]
    estadoNivelCalidad=['estado_nivel_calidad_lubricante_transmision'].iloc[0]
    estadoNivelCalidadObservacion=['estado_nivel_calidad_lubricante_transmision_observacion'].iloc[0]
    estadoNivelCalidadDiferencial=['estado_nivel_calidad_lubricante_diferencial'].iloc[0]
    estadoNivelCalidadDiferencialObservacion=['estado_nivel_calidad_lubricante_diferencial_observacion'].iloc[0]
    cubrepolvosFlechas = ['cubrepolvos_flechas'].iloc[0]
    cubrepolvosFlechasObservacion=['cubrepolvos_flechas_observacion'].iloc[0]
    componentesDireccion=['componentes_direccion'].iloc[0]
    componentesDireccionObservacion=['componentes_direccion_observacion'].iloc[0]
    componentesSuspension=['componentes_suspesion'].iloc[0]
    componentesSuspensionObservacion=['componentes_suspesion_observacion'].iloc[0]
    sistemasEscape=['sistema_escape_completo'].iloc[0]
    sistemasEscapeObservacion=['sistema_escape_completo_observacion'].iloc[0]
    sistemaAlimentacion=['sistema_alimentacion_combustible'].iloc[0]
    sistemaAlimentacionObservacion=['sistema_alimentacion_combustible_observacion'].iloc[0]
    filtroCombustible=['filtro_combustible'].iloc[0]
    filtroCombustibleObservacion=['filtro_combustible_observacion'].iloc[0]
    controlFugasDireccion=['control_fugas_direccion_hidraulica'].iloc[0]
    controlFugasDireccionObservacion=['control_fugas_direccion_hidraulica_observacion'].iloc[0]
    otroAltura=['otros_altura_total'].iloc[0]
    otrosAlturaObservacion=['otros_altura_total_observacion'].iloc[0]
    rodamientoMazas=['rodamiento_mazas_rueda'].iloc[0]
    rodamientoMAzasObservacion=['rodamiento_mazas_rueda_observacion'].iloc[0]
    holguraSuspension=['holgura_partes_suspension_rueda'].iloc[0]
    holguraSuspensionObservacion=['holgura_partes_suspension_rueda_observacion'].iloc[0]
    controlNeumaticos=['control_neumaticos_desgaste_presion'].iloc[0]
    controlNeumaticosObservacion=['control_neumaticos_desgaste_presion_observacion'].iloc[0]
    profundidad=['profundidad'].iloc[0]
    profundidadObservacion=['profundidad_observacion'].iloc[0]
    presion=['presion'].iloc[0]
    presionObservacion=['presion_observacion'].iloc[0]
    otrosAltura=['otros_altura_media'].iloc[0]
    otrosAlturaObservacion=['otros_altura_media_observacion'].iloc[0]
    nivelCalidadAceite=['nivel_calidad_aceite_motor'].iloc[0]
    nivelCalidadAceiteObservacion=['nivel_calidad_aceite_motor_observacion'].iloc[0]
    filtroAire=['filtro_aire'].iloc[0]
    filtroAireObservacion=['filtro_aire_observacion'].iloc[0]
    filtroPolen=['filtro_polen'].iloc[0]
    filtroPolenObservacion=['filtro_polen_observacion'].iloc[0]
    filtroPcv=['filtro_pcv'].iloc[0]
    filtroPcvObservacion=['filtro_pcv_observacion'].iloc[0]
    valvulaPcv=['valvula_pcv'].iloc[0]
    valvulaPcvObservacion=['valvula_pcv_observacion'].iloc[0]
    bujiasEncendido=['bujias_encendido'].iloc[0]
    bujiasEncendidoObservacion=['bujias_encendido_observacion'].iloc[0]
    cablesBujiasBobinas=['cables_bujias_bobinas_ignicion'].iloc[0]
    cablesBujiasBobinasObservacion=['cables_bujias_bobinas_ignicion_observacion'].iloc[0]
    nivelAnticongelante=['nivel_anticongenlante'].iloc[0]
    nivelAnticongelanteObservacion=['nivel_anticongenlante_observacion'].iloc[0]
    taponRadiador=['tapon_radiador'].iloc[0]
    taponRadiadorObservacion=['tapon_radiador_observacion'].iloc[0]
    manguerasSistema=['mangueras_sistema'].iloc[0]
    manguerasSistemaObservacion=['mangueras_sistema_observacion'].iloc[0]
    desempeñoVentilador=['desempeño_ventilador'].iloc[0]
    desempeñoVentiladorObservacion=['desempeño_ventilador_observacion'].iloc[0]
    calidadLiquidoLimpia=['calidad_liquido_limpiaparabrisas'].iloc[0]
    calidadLiquidoLimpiaObservacion=['calidad_liquido_limpiaparabrisas_observacion'].iloc[0]
    calidadAceiteDireccion=['calidad_aceite_direccion_hidraulica'].iloc[0]
    calidadAceiteDireccionObservacion=['calidad_aceite_direccion_hidraulica_observacion'].iloc[0]
    calidadAceiteTransmision=['calidad_aceite_transmision_bayoneta'].iloc[0]
    calidadAceiteTransmisionObservacion=['calidad_aceite_transmision_bayoneta_observacion'].iloc[0]
    liquidoBateriaCondiciones=['liquido_bateria_condiciones'].iloc[0]
    liquidoBateriaCondicionesObservacion=['liquido_bateria_condiciones_observacion'].iloc[0]
    bandasPoly=['bandas_poly_v'].iloc[0]
    bandasPolyObservacion=['bandas_poly_v_observacion'].iloc[0]
    poleasBanda=['poleas_banda'].iloc[0]
    poleasBandaObservacion=['poleas_banda_observacion'].iloc[0]
    bandaTiempo=['banda_tiempo'].iloc[0]
    bandaTiempoObservacion=['banda_tiempo_observacion'].iloc[0]
    otrosHabitaculo=['otros_habitaculo_motor'].iloc[0]
    otrosHabitaculoObservacion=['otros_habitaculo_motor_observacion'].iloc[0]
    resetIntervalo=['reset_intervalo_servicio'].iloc[0]
    resetIntervaloObservacion=['reset_intervalo_servicio_observacion'].iloc[0]
    ajusteTornillosNeumaticos=['ajuste_tornillos_neumaticos_torquimetro'].iloc[0]
    ajusteTornillosNeumaticosObservacion=['ajuste_tornillos_neumaticos_torquimetro_observacion'].iloc[0]
    limpiarLubricarPuertas=['limpiar_libricar_puertas_cerraduras'].iloc[0]
    limpiarLubricarPuertasObservacion=['limpiar_libricar_puertas_cerraduras_observacion'].iloc[0]
    completarPlanMantenimiento=['completar_plan_mantenimiento'].iloc[0]
    completarPlanMantenimientoObservacion=['completar_plan_mantenimiento_observacion'].iloc[0]
    NumeroSerie=['NumeroSerie'].iloc[0]

    htmlstring = r"""html:"<!DOCTYPE html> <html lang="es"> <head>
      <meta charset="UTF-8">
      <meta name="viewport" content="width=device-width, initial-scale=1.0">
      <title>Check List General</title>
      <style>
        body {
            font-family: Arial, sans-serif;
        }
        h1, h2 {
            text-align: center;
        }
        table {
            width: 100%;
            border-collapse: collapse;
            margin-bottom: 20px;
        }
        td, th {
            padding: 8px;
            text-align: left;
            border: 1px solid #ddd;
        }
        th {
            background-color: #f2f2f2;
        }
        .highlight {
            background-color: #f9e44c;
        }
        .footer {
            text-align: center;
            margin-top: 20px;
        }
     </style>
      </head>
    <body>"""
    htmlstring=f"""
      <h1>CHECK LIST GENERAL</h1>
      <h2>Aplicable para todos los servicios</h2>
    
      <table>
        <tr>
            <th>No. de Orden</th>
            <td>{idOrden}</td>
            <th>No. Folio</th>
            <td>484</td>
        </tr>
        <tr>
            <th>No. Serie del Vehículo</th>
            <td colspan="3">{NumeroSerie}</td>
        </tr>
        <tr>
            <th>Fecha</th>
            <td colspan="3">{fecha}</td>
        </tr>
      </table>

      <h3>Checklist</h3>

      <table>
        <thead>
            <tr>
                <th>Item</th>
                <th>OK</th>
                <th>NO OK</th>
                <th>Observaciones</th>
            </tr>
        </thead>
        <tbody>
            <tr class="highlight">
                <td>Vehículo en el piso</td>
                <td>OK</td>
                <td>NO OK</td>
                <td></td>
            </tr>
            <tr>
                <td>Lectura de códigos de falla. (Nº de averías en sus sistemas)</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{lecturaCodigoObservacion}</td>
            </tr>
            <tr>
                <td>Funcionamiento del servofreno</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{servofrenoObservacion}</td>
            </tr>
            <tr>
                <td>Funcionamiento del pedal de freno</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{pedalFrenoObservacion}</td>
            </tr>
            <tr>
                <td>Funcionamiento del pedal de estacionamiento</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{pedalEstacionamientoObservacion}</td>
            </tr>
            <tr>
                <td>Cuadro de instrumentos (funcionamiento, iluminación, testigos)</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{cuadroInstrumentosObservacion}</td>
            </tr>
            <tr>
                <td>Iluminación interior completa (cortesía, vanidad, guantera, etc.)</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{iluminacionInteriorObservacion}</td>
            </tr>
            <tr>
                <td>Iluminación exterior (cortes, bajas, altas, reversa, niebla, nivel)</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{iluminacionExternaObservacion}</td>
            </tr>
            <tr>
                <td>Sistema limpiaparabrisas completo, chorritos, plumas</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{limpiaparabrisasObservacion}</td>
            </tr>
            <tr>
                <td>Sistema limpia medallón completo, chorritos, plumas</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{limpiaMedallonObservacion}</td>
            </tr>
            <tr>
                <td>Sistema de frenos (desgaste y presión)</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{neumaticosFriccionObservacion}</td>
            </tr>
        </tbody>
    </table>

    <h3>Vehículo en altura total</h3>

    <table>
        <thead>
            <tr>
                <th>Item</th>
                <th>OK</th>
                <th>NO OK</th>
                <th>Observaciones</th>
            </tr>
        </thead>
        <tbody>
            <tr>
                <td>Estado de fugas de aceite de motor y transmisión</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{estadoFugasAceiteObservacion}</td>
            </tr>
            <tr>
                <td>Estado de nivel y calidad de lubricantes en transmisión</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{estadoNivelCalidadObservacion}</td>
            </tr>
            <tr>
                <td>Estado de nivel y calidad de lubricantes en diferencial</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{estadoNivelCalidadDiferencialObservacion}</td>
            </tr>
            <tr>
                <td>Estado de cubrepolvos de flechas homocinéticas</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{cubrepolvosFlechasObservacion}</td>
            </tr>
            <tr>
                <td>Estado de componentes de dirección</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{componentesDireccionObservacion}</td>
            </tr>
            <tr>
                <td>Estado de suspensión</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{componentesSuspensionObservacion}</td>
            </tr>
            <tr>
                <td>Estado del sistema de alimentación de combustible (fugas)</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{sistemaAlimentacionObservacion}</td>
            </tr>
            <tr>
                <td>Control de fugas de dirección hidráulica</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{controlFugasDireccionObservacion}</td>
            </tr>
        </tbody>
    </table>

    <h3>Vehículo en altura media</h3>

    <table>
        <thead>
            <tr>
                <th>Item</th>
                <th>OK</th>
                <th>NO OK</th>
                <th>Observaciones</th>
            </tr>
        </thead>
        <tbody>
            <tr>
                <td>Control de rodamiento y mazas de rueda</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{rodamientoMAzasObservacion}</td>
            </tr>
            <tr>
                <td>Holgura/juego en partes de suspensión y dirección</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{holguraSuspensionObservacion}</td>
            </tr>
            <tr>
                <td>Profundidad (D.D)</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{profundidadObservacion}</td>
            </tr>
            <tr>
                <td>Presión (D.D)</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{presionObservacion}</td>
            </tr>
        </tbody>
    </table>

    <h3>Control de Habitáculo de motor</h3>

    <table>
        <thead>
            <tr>
                <th>Item</th>
                <th>OK</th>
                <th>NO OK</th>
                <th>Observaciones</th>
            </tr>
        </thead>
        <tbody>
            <tr>
                <td>Nivel y calidad de aceite de motor</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{nivelCalidadAceiteObservacion}</td>
            </tr>
            <tr>
                <td>Filtro de aire</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{filtroAireObservacion}</td>
            </tr>
        </tbody>
    </table>

    <h3>Revisión Final</h3>

    <table>
        <thead>
            <tr>
                <th>Item</th>
                <th>OK</th>
                <th>NO OK</th>
                <th>Observaciones</th>
            </tr>
        </thead>
        <tbody>
            <tr>
                <td>Reset de intervalo de servicio</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{resetIntervaloObservacion}</td>
            </tr>
            <tr>
                <td>Ajuste y verificación de alineación</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{ajusteTornillosNeumaticosObservacion}</td>
            </tr>
            <tr>
                <td>Ajuste de frenos</td>
                <td>OK</td>
                <td>NO OK</td>
                <td></td>
            </tr>
            <tr>
                <td>Revisión de luces</td>
                <td>OK</td>
                <td>NO OK</td>
                <td></td>
            </tr>
        </tbody>
    </table>

    <div class="footer">
        <p>Firma Cliente: ___________________________</p>
        <p>Firma Responsable: ___________________________</p>
    </div>
    </body>
    </html>"""
    img = "\\img1.jpg"
    pdf_path = "example.pdf"
          # Rutas y configuraciones para Linux
    path_wkhtmltopdf = '/usr/local/bin/wkhtmltopdf'
    config = pdfkit.configuration(wkhtmltopdf=path_wkhtmltopdf)
    pdfkit.from_string(htmlstring, 'reporte.pdf', configuration=config)
    return JSONResponse(content={"message": "PDF creado exitosamente"}, status_code=200) 
    return JSONResponse(content={"error": str(e)}, status_code=500)

@app.get(
        path="/api/historicocheck",
        name='Obtener clientes',
        tags=['Historico'],
        description='Método para obtener la informacion de todos los clientes',
        response_model=List[GetCliente]
)
def getclientes(busqueda = ""):
    query = f"exec Clientes.clienteinsupdel @Accion = 2,@ParametroBusqueda = '{busqueda}' "
    roles_df = pd.read_sql(query, engine)
    resultado = roles_df.to_dict(orient="records")
    return JSONResponse(status_code=200,content=resultado)

@app.post(
        path="/api/ordenservice",
        name='Insertar orden de servicio',
        tags=['Orden'],
        description='Método para insertar la orden de servicio',
        response_model=OrdenService
)
def saveordenservice(payload: OrdenService):
    query = f"""EXEC InsertarOrdenServicio @idCliente = {payload.IdCliente} , \
        @idEmpleado = '{payload.IdEmpleado}' """
    print(query)

    with engine.begin() as conn:
        conn.execution_options(autocommit = True)
        roles_df = pd.read_sql(query, conn)
    dumpp = ResponseModel(id_resultado=1,respuesta="La orden se ha guardado de manera correcta")
    dict = dumpp.model_dump()
    return JSONResponse(status_code=200, content=dict)

if __name__ == '__main__':
    app.run()
    